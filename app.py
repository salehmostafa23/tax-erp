import streamlit as st
import pandas as pd
import json
import os
import uuid
import shutil
import math
import hashlib
import requests as http_requests
from datetime import datetime, timedelta
from io import BytesIO
import base64
import time
import openpyxl
from github_storage import gh_read, gh_write

st.set_page_config(page_title="Tax Management System", page_icon="🏢", layout="wide", initial_sidebar_state="expanded")

# ====================== DIRS ======================
DATA_DIR = os.path.dirname(os.path.abspath(__file__))

# ====================== USERS ======================
USERS_FILE=os.path.join(DATA_DIR,"users.json")
ALL_PAGES=["🏠 الرئيسية","📋 نموذج 41","💰 القيمة المضافة","🛒 فواتير الماركت","📄 Portal الفواتير الإلكترونية","🔍 الاستعلام عن ممول","🏷️ الاستعلام عن الأكواد"]
ADMIN_PAGE="👥 إدارة المستخدمين"

def _hash_pw(pw,salt="tax_erp_salt_2024"):
    return hashlib.sha256(f"{salt}{pw}".encode()).hexdigest()

def load_users():
    gh=gh_read("users.json")
    if gh is not None:
        with open(USERS_FILE,'w',encoding='utf-8') as f: json.dump(gh,f,ensure_ascii=False,indent=2,default=str)
        return gh
    if os.path.exists(USERS_FILE):
        try:
            with open(USERS_FILE,'r',encoding='utf-8') as f: return json.load(f)
        except: pass
    default=[{"username":"admin","password":_hash_pw("admin123"),"display_name":"المدير","role":"admin","permissions":ALL_PAGES+[ADMIN_PAGE],"created_at":datetime.now().isoformat()}]
    save_users(default)
    return default

def save_users(data):
    with open(USERS_FILE,'w',encoding='utf-8') as f: json.dump(data,f,ensure_ascii=False,indent=2,default=str)
    gh_write("users.json",data)

def authenticate_user(username,password):
    users=load_users()
    pw_hash=_hash_pw(password)
    for u in users:
        if u['username']==username and u['password']==pw_hash:
            return u
    return None

def get_current_user():
    return st.session_state.get('current_user',None)

def user_has_permission(page):
    u=get_current_user()
    if not u: return False
    if u.get('role')=='admin': return True
    return page in u.get('permissions',[])

# ====================== ETA API ======================
ETA_IDENTITY_URL = "https://id.eta.gov.eg/connect/token"
ETA_API_BASE = "https://api.invoicing.eta.gov.eg"
ETA_DOC_TYPE_MAP = {"i":"فاتورة بيع","c":"إشعار دائن","d":"إشعار مدين","ii":"فاتورة استيراد","ei":"فاتورة تصدير","ec":"إشعار دائن تصدير","ed":"إشعار مدين تصدير"}
ETA_STATUS_MAP = {"Valid":"مقبولة","Invalid":"مرفوضة","Submitted":"مرسلة","Rejected":"مرفوضة","Cancelled":"ملغاة"}

def eta_login(client_id, client_secret):
    cred = f"{client_id}:{client_secret}"
    encoded = base64.b64encode(cred.encode()).decode()
    headers = {"Authorization": f"Basic {encoded}", "Content-Type": "application/x-www-form-urlencoded"}
    data = {"grant_type": "client_credentials"}
    r = http_requests.post(ETA_IDENTITY_URL, headers=headers, data=data, timeout=30, verify=False)
    if r.status_code == 200:
        return r.json().get("access_token"), None
    try:
        err = r.json()
        return None, err.get("error_description", err.get("error", r.text[:200]))
    except:
        return None, f"HTTP {r.status_code}: {r.text[:200]}"

def eta_search_docs(token, direction, date_from, date_to, page_size=100):
    url = f"{ETA_API_BASE}/api/v1.0/documents/search"
    headers = {"Authorization": f"Bearer {token}"}
    params = {
        "issueDateFrom": date_from.strftime("%Y-%m-%dT00:00:00"),
        "issueDateTo": date_to.strftime("%Y-%m-%dT23:59:59"),
        "direction": direction,
        "pageSize": page_size
    }
    all_docs = []
    continuation = None
    while True:
        if continuation:
            params["continuationToken"] = continuation
        r = http_requests.get(url, headers=headers, params=params, timeout=30, verify=False)
        if r.status_code != 200:
            try:
                err = r.json()
                return None, err.get("error_description", err.get("error", f"HTTP {r.status_code}"))
            except:
                return None, f"HTTP {r.status_code}: {r.text[:200]}"
        body = r.json()
        docs = body.get("result", [])
        all_docs.extend(docs)
        meta = body.get("metadata", {})
        continuation = meta.get("continuationToken", "")
        if not continuation or continuation == "EndofResultSet":
            break
    return all_docs, None

def eta_doc_to_record(doc, direction):
    uuid_val = doc.get("uuid", "")
    internal_id = doc.get("internalId", "")
    type_name = doc.get("typeName", "i")
    status_raw = doc.get("status", "")
    status = ETA_STATUS_MAP.get(status_raw, status_raw)
    issue_date = doc.get("dateTimeIssued", "")
    submit_date = doc.get("dateTimeReceived", "")
    issuer_id = doc.get("issuerId", "")
    issuer_name = doc.get("issuerName", "")
    receiver_id = doc.get("receiverId", "")
    receiver_name = doc.get("receiverName", "")
    total_sales = float(doc.get("totalSales", 0) or 0)
    total_discount = float(doc.get("totalDiscount", 0) or 0)
    net_amount = float(doc.get("netAmount", 0) or 0)
    total = float(doc.get("total", 0) or 0)
    tax_total = round(total - net_amount, 2) if total and net_amount else 0
    if direction == "Sent":
        period = issue_date[:7].replace("-", "/") if issue_date else ""
        counterparty = receiver_name
        counterparty_id = receiver_id
    else:
        period = issue_date[:7].replace("-", "/") if issue_date else ""
        counterparty = issuer_name
        counterparty_id = issuer_id
    rec = {
        "UUID": uuid_val, "internalId": internal_id,
        "نوع الفاتورة": ETA_DOC_TYPE_MAP.get(type_name, type_name),
        "الحالة": status,
        "تاريخ الإصدار": issue_date[:10] if issue_date else "",
        "تاريخ الإرسال": submit_date[:10] if submit_date else "",
        "رقم التسجيل (المصدر)": issuer_id, "اسم المصدر": issuer_name,
        "رقم التسجيل (المستلم)": receiver_id, "اسم المستلم": receiver_name,
        "الطرف الآخر": counterparty, "رقم التسجيل (الطرف الآخر)": counterparty_id,
        "إجمالي المبيعات (قبل الخصم)": total_sales, "الخصم": total_discount,
        "ضريبة القيمة المضافة": tax_total,
        "الإجمالي (بعد الضريبة)": total
    }
    meta = {"uuid": uuid_val, "upload_date": submit_date or datetime.now().isoformat(),
        "period": period, "invoice_type": ETA_DOC_TYPE_MAP.get(type_name, type_name),
        "status": status, "file_name": f"ETA_{uuid_val[:12]}.json",
        "source": "eta_api", "records": [], "records_count": 1}
    return rec, meta

def eta_get_document_pdf(token, uuid_val):
    url = f"{ETA_API_BASE}/api/v1.0/documents/{uuid_val}/pdf"
    headers = {"Authorization": f"Bearer {token}"}
    r = http_requests.get(url, headers=headers, timeout=30, verify=False, stream=True)
    if r.status_code == 200:
        buf = BytesIO()
        for chunk in r.iter_content(chunk_size=8192):
            buf.write(chunk)
        buf.seek(0)
        return buf, None
    return None, f"HTTP {r.status_code}"

CODES_DB_FILE=os.path.join(DATA_DIR,"codes_database.json")
def load_codes_db():
    gh=gh_read("codes_database.json")
    if gh is not None:
        with open(CODES_DB_FILE,'w',encoding='utf-8') as f: json.dump(gh,f,ensure_ascii=False,indent=2,default=str)
        return gh
    if os.path.exists(CODES_DB_FILE):
        try:
            with open(CODES_DB_FILE,'r',encoding='utf-8') as f: return json.load(f)
        except: pass
    return []
def save_codes_db(data):
    with open(CODES_DB_FILE,'w',encoding='utf-8') as f: json.dump(data,f,ensure_ascii=False,indent=2,default=str)
    gh_write("codes_database.json",data)

def eta_get_document_details(token, uuid_val):
    url=f"{ETA_API_BASE}/api/v1.0/documents/{uuid_val}/details"
    headers={"Authorization":f"Bearer {token}"}
    r=http_requests.get(url,headers=headers,timeout=30,verify=False)
    if r.status_code==200:
        return r.json(),None
    return None,f"HTTP {r.status_code}"

def _fix_vat_in_records(records):
    fixed=[]
    for r in records:
        r=dict(r)
        total=_sf(r.get('الإجمالي (بعد الضريبة)',0))
        net=_sf(r.get('الصافي (قبل الضريبة)',0))
        current_vat=_sf(r.get('ضريبة القيمة المضافة',0))
        if not current_vat and total and net and total>net:
            r['ضريبة القيمة المضافة']=round(total-net,2)
        elif not current_vat and total and net and abs(total-net)<0.01:
            sales=_sf(r.get('إجمالي المبيعات (قبل الخصم)',0))
            disc=_sf(r.get('الخصم',0))
            if sales and disc and sales>disc:
                r['ضريبة القيمة المضافة']=round(total-(sales-disc),2)
        fixed.append(r)
    return fixed

def _generate_pdf_for_records(records, title="فواتير"):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    buf = BytesIO()
    page_w, page_h = landscape(A4)
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), rightMargin=1*cm, leftMargin=1*cm, topMargin=1*cm, bottomMargin=1*cm)
    styles = getSampleStyleSheet()
    try:
        pdfmetrics.registerFont(TTFont('Arabic', 'C:/Windows/Fonts/arial.ttf'))
        arabic_style = ParagraphStyle('Arabic', parent=styles['Normal'], fontName='Arabic', fontSize=7, leading=10, alignment=1)
        title_style = ParagraphStyle('TitleA', parent=styles['Title'], fontName='Arabic', fontSize=14, alignment=1)
    except:
        arabic_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=7, leading=10, alignment=1)
        title_style = styles['Title']
    elements = []
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 0.3*cm))
    key_fields = ['نوع الفاتورة','الحالة','تاريخ الإصدار','الطرف الآخر','رقم التسجيل (الطرف الآخر)','إجمالي المبيعات (قبل الخصم)','الخصم','الصافي ( trước الضريبة)','ضريبة القيمة المضافة','الإجمالي (بعد الضريبة)','UUID']
    for idx, rec in enumerate(records):
        rec_status = rec.get('الحالة', '')
        elements.append(Paragraph(f"فاتورة #{idx+1} — {rec.get('الطرف الآخر', '-')} — الإجمالي: {rec.get('الإجمالي (بعد الضريبة)', 0)} — الحالة: {rec_status}", arabic_style))
        elements.append(Spacer(1, 0.15*cm))
        pairs = []
        for k in rec.keys():
            v = rec.get(k, '')
            pairs.append([Paragraph(str(k), arabic_style), Paragraph(str(v)[:60], arabic_style)])
        if pairs:
            usable_w = page_w - 2*cm
            t = Table(pairs, colWidths=[usable_w*0.4, usable_w*0.6])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#6c5ce7')),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 7),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ROWBACKGROUNDS', (1, 0), (1, -1), [colors.white, colors.HexColor('#f0f0f8')]),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            elements.append(t)
        elements.append(Spacer(1, 0.4*cm))
    doc.build(elements)
    buf.seek(0)
    return buf

# ====================== LOGIN ======================
if 'current_user' not in st.session_state:
    st.session_state['current_user']=None

if not st.session_state['current_user']:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=Cairo:wght@300;400;500;600;700;800&display=swap');
    :root{--bg:#0a0a15;--surface:rgba(22,22,40,0.7);--surface2:#1e1e38;--border:rgba(255,255,255,0.06);--text:#eaeaf2;--text2:#7878a0;--accent:#6c5ce7;--accent2:#a29bfe;}
    html,body,[class*="css"]{font-family:'Inter','Cairo',sans-serif!important;}
    #MainMenu,footer,header,.stDeployButton{visibility:hidden!important;}
    div[data-testid="stToolbar"]{display:none!important;}
    .stApp{background:linear-gradient(135deg,#08081a 0%,#0d0d22 50%,#0a0a18 100%)!important;}
    .login-box{max-width:320px;margin:6rem auto 0;padding:1.5rem 1.5rem;border-radius:16px;background:rgba(22,22,40,0.7);border:1px solid rgba(255,255,255,0.06);backdrop-filter:blur(20px);box-shadow:0 20px 60px rgba(0,0,0,.5);position:relative;overflow:hidden;}
    .login-box::before{content:'';position:absolute;top:-50%;left:-50%;width:200%;height:200%;background:radial-gradient(circle at center,rgba(108,92,231,0.06),transparent 50%);pointer-events:none;}
    .login-icon{width:50px;height:50px;margin:0 auto .8rem;border-radius:14px;background:linear-gradient(135deg,#6c5ce7 0%,#a29bfe 50%,#00cec9 100%);display:flex;align-items:center;justify-content:center;font-size:1.5rem;box-shadow:0 8px 30px rgba(108,92,231,0.45);}
    .login-title{margin:0 0 .2rem;color:#fff;font-size:1.1rem;font-weight:800;text-align:center;letter-spacing:.5px;}
    .login-sub{margin:0 0 1rem;color:rgba(255,255,255,.3);font-size:.65rem;text-align:center;letter-spacing:1px;}
    .stTextInput>div>div>input{background:rgba(30,30,56,0.8)!important;border:1px solid rgba(255,255,255,0.08)!important;border-radius:12px!important;color:#fff!important;padding:.7rem 1rem!important;}
    .stTextInput>div>div>input:focus{border-color:rgba(108,92,231,0.5)!important;box-shadow:0 0 0 3px rgba(108,92,231,0.1)!important;}
    .stButton>button[kind="primary"]{background:linear-gradient(135deg,#6c5ce7 0%,#a29bfe 100%)!important;border:none!important;border-radius:12px!important;font-family:'Cairo',sans-serif!important;font-weight:700!important;color:#fff!important;padding:.6rem 0!important;width:100%!important;box-shadow:0 4px 20px rgba(108,92,231,0.4)!important;transition:all .3s!important;font-size:.95rem!important;}
    .stButton>button[kind="primary"]:hover{transform:translateY(-2px)!important;box-shadow:0 8px 30px rgba(108,92,231,0.5)!important;}
    .login-err{background:rgba(255,107,107,.1);border:1px solid rgba(255,107,107,.2);border-radius:10px;padding:.5rem 1rem;color:#ff6b6b;font-size:.8rem;text-align:center;margin-top:.5rem;}
    </style>
    """, unsafe_allow_html=True)
    st.markdown('<div class="login-box"><div class="login-icon">🏢</div><h2 class="login-title">Tax Management System</h2><p class="login-sub">نظام إدارة الضرائب المتكامل</p></div>', unsafe_allow_html=True)
    with st.form("login_form",clear_on_submit=False):
        u=st.text_input("اسم المستخدم",key="login_user",placeholder="Username")
        p=st.text_input("كلمة المرور",key="login_pass",type="password",placeholder="Password")
        submitted=st.form_submit_button("تسجيل الدخول",type="primary",use_container_width=True)
        if submitted:
            if not u or not p:
                st.markdown('<div class="login-err">أدخل اسم المستخدم وكلمة المرور</div>', unsafe_allow_html=True)
            else:
                user=authenticate_user(u.strip(),p)
                if user:
                    st.session_state['current_user']=user
                    st.rerun()
                else:
                    st.markdown('<div class="login-err">بيانات الدخول غير صحيحة</div>', unsafe_allow_html=True)
    st.markdown('<p style="text-align:center;color:rgba(255,255,255,.2);font-size:.6rem;margin-top:2rem;font-family:Cairo,sans-serif;">جميع الحقوق محفوظة © تصميم محاسب / صالح مصطفى</p>', unsafe_allow_html=True)
    st.stop()

# ====================== CSS ======================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&family=Cairo:wght@300;400;500;600;700;800&display=swap');
:root{--bg:#0a0a15;--surface:rgba(22,22,40,0.7);--surface2:#1e1e38;--border:rgba(255,255,255,0.06);--text:#eaeaf2;--text2:#7878a0;--accent:#6c5ce7;--accent2:#a29bfe;--cyan:#00cec9;--green:#00b894;--orange:#fdcb6e;--red:#ff6b6b;--pink:#fd79a8;--blue:#74b9ff;}
html,body,[class*="css"]{font-family:'Inter','Cairo',sans-serif!important;direction:rtl;}
#MainMenu,footer,header,.stDeployButton{visibility:hidden!important;}
div[data-testid="stToolbar"]{display:none!important;}
.stApp{background:linear-gradient(135deg,#08081a 0%,#0d0d22 50%,#0a0a18 100%)!important;color:var(--text);}
.block-container{padding:1rem 2rem 2rem 2rem!important;max-width:100%!important;}

/* SIDEBAR */
section[data-testid="stSidebar"]{background:linear-gradient(180deg,#050510 0%,#0b0b24 40%,#080820 100%)!important;border-left:1px solid rgba(108,92,231,0.08)!important;box-shadow:8px 0 60px rgba(0,0,0,0.7)!important;position:relative!important;}
section[data-testid="stSidebar"][aria-expanded="false"]{position:relative!important;transform:none!important;margin-left:0!important;}
button[data-testid="stSidebarCollapseButton"],div[data-testid="stSidebarCollapseButton"],[data-testid="stSidebarCollapseButton"]{display:none!important;pointer-events:none!important;visibility:hidden!important;opacity:0!important;width:0!important;height:0!important;padding:0!important;margin:0!important;overflow:hidden!important;}
section[data-testid="stSidebar"]>div:first-child{padding-top:0!important;}
section[data-testid="stSidebar"] [data-testid="stSidebarContent"]{overflow-y:auto!important;max-height:calc(100vh - 2rem)!important;}
section[data-testid="stSidebar"] .stMarkdown p,section[data-testid="stSidebar"] .stMarkdown span,section[data-testid="stSidebar"] label,section[data-testid="stSidebar"] .stRadio>div>label{color:rgba(255,255,255,0.55)!important;font-size:.8rem!important;font-family:'Inter','Cairo',sans-serif!important;}
section[data-testid="stSidebar"] .stRadio>div>div>label::before{display:none!important;}
section[data-testid="stSidebar"] .stRadio>div>div>label::after{display:none!important;}
section[data-testid="stSidebar"] .stRadio>div{gap:0!important;}
section[data-testid="stSidebar"] .stRadio{gap:0!important;}
section[data-testid="stSidebar"] [data-testid="stRadio"]>div{gap:0!important;}
section[data-testid="stSidebar"] .stRadio>div>div>label{background:rgba(108,92,231,0.1)!important;border:1px solid rgba(108,92,231,0.2)!important;border-radius:8px!important;padding:.3rem .7rem!important;margin:2px 4px!important;display:flex!important;align-items:center!important;gap:.5rem!important;font-size:.75rem!important;}
section[data-testid="stSidebar"] .stRadio>div>div>label:hover{background:rgba(108,92,231,0.2)!important;border-color:rgba(108,92,231,0.4)!important;color:rgba(255,255,255,.9)!important;}
section[data-testid="stSidebar"] .stRadio>div>div:has(input:checked)>label{background:rgba(108,92,231,0.25)!important;border:1px solid rgba(108,92,231,0.6)!important;color:#fff!important;font-weight:700!important;box-shadow:0 0 12px rgba(108,92,231,0.15)!important;}
section[data-testid="stSidebar"] hr{border-color:rgba(108,92,231,.06)!important;}
section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p{margin:0!important;}

/* TOPBAR */
.erp-topbar{background:linear-gradient(135deg,rgba(108,92,231,0.07) 0%,rgba(0,206,201,0.04) 100%);border:1px solid rgba(108,92,231,0.1);border-radius:20px;padding:1.2rem 2rem;margin:0 0 1.5rem 0;display:flex;align-items:center;justify-content:space-between;backdrop-filter:blur(20px);position:relative;overflow:hidden;}
.erp-topbar::before{content:'';position:absolute;top:-50%;right:-10%;width:300px;height:300px;background:radial-gradient(circle,rgba(108,92,231,0.06) 0%,transparent 70%);border-radius:50%;}
.erp-topbar::after{content:'';position:absolute;bottom:-50%;left:-5%;width:200px;height:200px;background:radial-gradient(circle,rgba(0,206,201,0.05) 0%,transparent 70%);border-radius:50%;}
.erp-topbar h2{margin:0;font-size:1.3rem;font-weight:700;background:linear-gradient(135deg,#fff 0%,#a29bfe 100%);-webkit-background-clip:text;-webkit-text-fill-color:transparent;}
.erp-topbar p{margin:.2rem 0 0;font-size:.78rem;color:var(--text2);}
.erp-topbar-right{display:flex;align-items:center;gap:.8rem;z-index:1;}
.erp-topbar-right button:hover{background:rgba(108,92,231,.3)!important;border-color:rgba(108,92,231,.5)!important;transform:scale(1.05);}
.erp-badge{background:rgba(108,92,231,0.12);border:1px solid rgba(108,92,231,0.2);padding:.35rem .9rem;border-radius:20px;color:var(--accent2);font-size:.72rem;font-weight:600;}
.erp-time{color:var(--text2);font-size:.75rem;}

/* STAT CARDS */
.erp-stat{background:var(--surface);border:1px solid var(--border);border-radius:16px;padding:1.3rem;text-align:center;position:relative;overflow:hidden;transition:all .4s cubic-bezier(.4,0,.2,1);backdrop-filter:blur(10px);}
.erp-stat::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;border-radius:16px 16px 0 0;}
.erp-stat:hover{transform:translateY(-4px);box-shadow:0 20px 40px rgba(0,0,0,.3);border-color:rgba(255,255,255,.08);}
.erp-stat-label{font-size:.7rem;color:var(--text2);font-weight:500;letter-spacing:.5px;}
.erp-stat-value{font-size:1.8rem;font-weight:800;margin:.3rem 0;line-height:1;}
.erp-stat-sub{font-size:.65rem;color:var(--text2);opacity:.7;}
.s-blue::before{background:linear-gradient(90deg,#6c5ce7,#a29bfe)}.s-blue .erp-stat-value{color:#a29bfe}
.s-cyan::before{background:linear-gradient(90deg,#00cec9,#55efc4)}.s-cyan .erp-stat-value{color:#55efc4}
.s-orange::before{background:linear-gradient(90deg,#e17055,#fdcb6e)}.s-orange .erp-stat-value{color:#fdcb6e}
.s-green::before{background:linear-gradient(90deg,#00b894,#55efc4)}.s-green .erp-stat-value{color:#55efc4}
.s-pink::before{background:linear-gradient(90deg,#e84393,#fd79a8)}.s-pink .erp-stat-value{color:#fd79a8}
.s-red::before{background:linear-gradient(90deg,#ff6b6b,#ff9f9f)}.s-red .erp-stat-value{color:#ff6b6b}

/* GLASS CARD */
.erp-card{background:rgba(22,22,40,0.5);border:1px solid rgba(255,255,255,0.05);border-radius:16px;padding:1.5rem;backdrop-filter:blur(20px);margin-bottom:1rem;transition:all .3s ease;}
.erp-card:hover{border-color:rgba(108,92,231,0.12);box-shadow:0 8px 32px rgba(0,0,0,.2);}

/* BATCH CARD — the premium one */
.erp-batch{background:var(--surface);border:1px solid var(--border);border-radius:20px;padding:0;overflow:hidden;transition:all .4s cubic-bezier(.4,0,.2,1);backdrop-filter:blur(10px);position:relative;}
.erp-batch::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;transition:all .4s;}
.erp-batch.b-f41::before{background:linear-gradient(90deg,#6c5ce7,#a29bfe);}
.erp-batch.b-vat::before{background:linear-gradient(90deg,#00cec9,#55efc4);}
.erp-batch:hover{transform:translateY(-6px) scale(1.01);box-shadow:0 25px 60px rgba(0,0,0,.35);border-color:rgba(108,92,231,.15);}
.erp-batch-head{padding:1.2rem 1.5rem .8rem;display:flex;align-items:center;gap:.8rem;}
.erp-batch-icon{width:46px;height:46px;border-radius:14px;display:flex;align-items:center;justify-content:center;font-size:1.4rem;flex-shrink:0;}
.erp-batch-title{font-size:.95rem;font-weight:700;color:var(--text);line-height:1.3;}
.erp-batch-sub{font-size:.7rem;color:var(--text2);margin-top:.1rem;}
.erp-batch-grid{display:grid;grid-template-columns:1fr 1fr;gap:.6rem 1.2rem;padding:.5rem 1.5rem 1rem;}
.erp-batch-k{font-size:.62rem;color:var(--text2);font-weight:500;text-transform:uppercase;letter-spacing:.5px;}
.erp-batch-v{font-size:.88rem;font-weight:700;color:var(--text);margin-top:.1rem;}
.erp-batch-foot{padding:.8rem 1.5rem;border-top:1px solid var(--border);display:flex;align-items:center;justify-content:space-between;}
.erp-batch-count{font-size:.75rem;color:var(--text2);}
.erp-batch-count span{color:var(--accent2);font-weight:700;}

/* DETAIL VIEW */
.erp-detail-header{display:flex;align-items:center;gap:1rem;margin-bottom:1.5rem;}
.erp-back{display:inline-flex;align-items:center;gap:.4rem;padding:.5rem 1.2rem;border-radius:12px;background:var(--surface);border:1px solid var(--border);color:var(--text);font-size:.82rem;font-weight:600;cursor:pointer;transition:all .3s;text-decoration:none;}
.erp-back:hover{border-color:rgba(108,92,231,.3);background:rgba(108,92,231,.08);transform:translateX(-3px);}
.erp-detail-title{font-size:1.1rem;font-weight:700;color:var(--text);margin:0;}
.erp-meta-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(180px,1fr));gap:.8rem;margin-bottom:1.5rem;}
.erp-meta-item{background:var(--surface);border:1px solid var(--border);border-radius:14px;padding:1rem;backdrop-filter:blur(10px);transition:all .3s;}
.erp-meta-item:hover{border-color:rgba(108,92,231,.15);}
.erp-meta-k{font-size:.62rem;color:var(--text2);font-weight:600;text-transform:uppercase;letter-spacing:.8px;margin-bottom:.3rem;}
.erp-meta-v{font-size:1rem;font-weight:700;color:var(--text);}

/* SECTION */
.erp-section{display:flex;align-items:center;gap:.6rem;margin:1.5rem 0 1rem;}
.erp-section-dot{width:8px;height:8px;border-radius:50%;background:var(--accent);box-shadow:0 0 12px rgba(108,92,231,.5);}
.erp-section h3{margin:0;font-size:.95rem;font-weight:700;color:var(--text);}

/* INFO */
.erp-info{background:rgba(108,92,231,.06);border:1px solid rgba(108,92,231,.12);padding:.8rem 1.2rem;border-radius:12px;font-size:.8rem;color:var(--text2);line-height:1.7;}

/* FORM */
.stTextInput>div>div>input{background:var(--surface2)!important;border:1px solid var(--border)!important;border-radius:10px!important;color:var(--text)!important;}
.stTextInput>div>div>input:focus{border-color:rgba(108,92,231,.4)!important;box-shadow:0 0 0 3px rgba(108,92,231,.08)!important;}
.stSelectbox>div>div{background:var(--surface2)!important;border:1px solid var(--border)!important;border-radius:10px!important;}
.stDateInput>div>div>div{background:var(--surface2)!important;border:1px solid var(--border)!important;border-radius:10px!important;}
.stFileUploader{border:2px dashed rgba(108,92,231,.25)!important;border-radius:12px!important;padding:1rem!important;background:rgba(108,92,231,.02)!important;}
.stFileUploader:hover{border-color:rgba(108,92,231,.4)!important;}

/* BUTTONS */
.stButton>button[kind="primary"],.stDownloadButton>button{background:linear-gradient(135deg,#6c5ce7 0%,#a29bfe 100%)!important;border:none!important;border-radius:10px!important;font-family:'Cairo',sans-serif!important;font-weight:600!important;color:#fff!important;padding:.5rem 2rem!important;box-shadow:0 4px 15px rgba(108,92,231,.35)!important;transition:all .3s!important;}
.stButton>button[kind="primary"]:hover,.stDownloadButton>button:hover{transform:translateY(-2px)!important;box-shadow:0 8px 25px rgba(108,92,231,.45)!important;}
.stButton>button{background:var(--surface2)!important;border:1px solid var(--border)!important;border-radius:10px!important;color:var(--text)!important;font-family:'Cairo',sans-serif!important;}
.stButton>button:hover{border-color:rgba(108,92,231,.25)!important;background:rgba(108,92,231,.08)!important;}

/* TABS */
.stTabs [data-baseweb="tab-list"]{gap:4px;background:rgba(22,22,40,.4)!important;padding:4px!important;border-radius:12px!important;border:1px solid var(--border)!important;}
.stTabs [data-baseweb="tab"]{padding:10px 20px!important;border-radius:10px!important;font-weight:600!important;font-size:.82rem!important;color:var(--text2)!important;background:transparent!important;}
.stTabs [aria-selected="true"]{background:var(--surface)!important;color:var(--accent2)!important;box-shadow:0 2px 10px rgba(0,0,0,.2)!important;border:1px solid rgba(108,92,231,.15)!important;}

.stDataFrame{border-radius:12px!important;overflow:hidden!important;border:1px solid var(--border)!important;}
.stAlert{border-radius:10px!important;}

.erp-empty{text-align:center;padding:4rem 2rem;background:var(--surface);border:1px solid var(--border);border-radius:20px;}
.erp-empty-icon{width:80px;height:80px;margin:0 auto 1.5rem;border-radius:20px;background:rgba(108,92,231,.08);display:flex;align-items:center;justify-content:center;font-size:2.5rem;}
.erp-empty h3{color:var(--text);margin:0 0 .5rem;font-size:1.1rem;}
.erp-empty p{color:var(--text2);margin:0;font-size:.85rem;}

::-webkit-scrollbar{width:5px}::-webkit-scrollbar-track{background:transparent}::-webkit-scrollbar-thumb{background:rgba(108,92,231,.25);border-radius:5px}
</style>
""", unsafe_allow_html=True)

# ====================== SIDEBAR ======================
with st.sidebar:
    st.markdown("""
    <div style="text-align:center;padding:1.2rem 1rem .8rem;position:relative;">
        <div style="position:absolute;top:0;left:0;right:0;height:100%;background:radial-gradient(ellipse at center top,rgba(108,92,231,0.08),transparent 70%);pointer-events:none;"></div>
        <div style="width:56px;height:56px;border-radius:18px;background:linear-gradient(135deg,#6c5ce7 0%,#a29bfe 50%,#00cec9 100%);display:inline-flex;align-items:center;justify-content:center;font-size:1.6rem;margin-bottom:.7rem;box-shadow:0 8px 30px rgba(108,92,231,0.45);position:relative;">
            <div style="position:absolute;inset:-2px;border-radius:20px;background:linear-gradient(135deg,#6c5ce7,#a29bfe,#00cec9);z-index:-1;opacity:.3;filter:blur(8px);"></div>
            🏢
        </div>
        <h3 style="margin:0;color:#fff;font-size:1.05rem;font-weight:800;letter-spacing:.5px;">Tax Management System</h3>
        <p style="margin:.35rem 0 0;color:rgba(255,255,255,.28);font-size:.6rem;font-weight:500;letter-spacing:1px;">نظام إدارة الضرائب المتكامل</p>
        <div style="width:40px;height:2px;background:linear-gradient(90deg,transparent,rgba(108,92,231,.4),transparent);margin:.8rem auto 0;border-radius:2px;"></div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("<p style='color:rgba(255,255,255,.18);font-size:.55rem;font-weight:700;letter-spacing:3px;padding:.2rem .8rem;margin:.2rem 0 .3rem;text-transform:uppercase;'>القائمة</p>", unsafe_allow_html=True)
    cu=get_current_user()
    try:
        if cu and cu.get('role')=='admin':
            cur_perms=cu.get('permissions',[])
            if ADMIN_PAGE not in cur_perms: cur_perms.append(ADMIN_PAGE)
            missing=[p for p in ALL_PAGES if p not in cur_perms]
            if missing:
                cu['permissions']=cur_perms+missing
                users=load_users()
                for u in users:
                    if u.get('username')==cu.get('username'):
                        u['permissions']=cu['permissions']
                save_users(users)
    except Exception:
        pass
    nav_pages=[p for p in ALL_PAGES if p in cu.get('permissions',[]) or cu.get('role')=='admin']
    if cu.get('role')=='admin': nav_pages.append(ADMIN_PAGE)
    if 'selected_page' not in st.session_state or st.session_state.get('selected_page') not in nav_pages:
        st.session_state['selected_page']=nav_pages[0]
    for _np in nav_pages:
        _sel=(st.session_state['selected_page']==_np)
        if st.button(_np,key=f"navbtn_{_np}",use_container_width=True,type="primary" if _sel else "secondary"):
            st.session_state['selected_page']=_np
            st.rerun()
    page=st.session_state['selected_page']
    st.markdown("<div style='height:1px;background:linear-gradient(90deg,transparent,rgba(108,92,231,.1),transparent);margin:1.2rem .8rem;'></div>", unsafe_allow_html=True)
    st.markdown(f"""<div style="padding:.7rem 1rem;border-radius:14px;background:rgba(108,92,231,.04);border:1px solid rgba(108,92,231,.06);margin:0 .5rem;text-align:center;">
        <p style="color:rgba(255,255,255,.6);font-size:.65rem;margin:0 0 .2rem;">المستخدم: <strong style="color:#a29bfe;">{cu.get('display_name','')}</strong></p>
        <p style="color:rgba(255,255,255,.25);font-size:.55rem;margin:0;">{cu.get('role','user')}</p></div>""", unsafe_allow_html=True)
    if st.button("🚪 خروج",key="logout_btn",use_container_width=True):
        st.session_state['current_user']=None
        for k in list(st.session_state.keys()):
            if k!='current_user': del st.session_state[k]
        st.rerun()
    st.markdown("<div style='height:1px;background:linear-gradient(90deg,transparent,rgba(108,92,231,.1),transparent);margin:1.2rem .8rem;'></div>", unsafe_allow_html=True)
    st.markdown("""<div style="padding:.8rem 1rem;border-radius:14px;background:linear-gradient(135deg,rgba(108,92,231,.04),rgba(0,206,201,.02));border:1px solid rgba(108,92,231,.06);margin:0 .5rem;text-align:center;">
        <div style="display:flex;align-items:center;justify-content:center;gap:.4rem;margin-bottom:.3rem;">
            <div style="width:5px;height:5px;border-radius:50%;background:#00b894;box-shadow:0 0 6px #00b894;"></div>
            <p style="color:rgba(255,255,255,.5);font-size:.62rem;margin:0;font-weight:500;">متصل</p>
        </div>
        <p style="color:rgba(255,255,255,.22);font-size:.55rem;margin:0;letter-spacing:1px;">v1.0.0 • Tax Management System</p></div>""", unsafe_allow_html=True)

# ====================== DATA ======================
FORM41_FILE = os.path.join(DATA_DIR, "form41_data.json")
VAT_FILE = os.path.join(DATA_DIR, "vat_data.json")
PORTAL_OUT_FILE = os.path.join(DATA_DIR, "portal_outgoing.json")
PORTAL_IN_FILE = os.path.join(DATA_DIR, "portal_incoming.json")

def _gh_key(f):
    return os.path.basename(f)

def save_data(f,d):
    with open(f,'w',encoding='utf-8') as fh: json.dump(d,fh,ensure_ascii=False,indent=2,default=str)
    gh_write(_gh_key(f),d)
def load_data(f):
    gh=gh_read(_gh_key(f))
    if gh is not None:
        with open(f,'w',encoding='utf-8') as fh: json.dump(gh,fh,ensure_ascii=False,indent=2,default=str)
        return gh
    if os.path.exists(f):
        try:
            with open(f,'r',encoding='utf-8') as fh: return json.load(fh)
        except: return []
    return []
def read_form41_excel(f):
    df=pd.read_excel(f,header=None,engine='openpyxl')
    target=['م','رقم التسجيل الضريبي','اسم الممول','تاريخ التعامل','طبيعة التعامل','القيمة الإجمالية للتعامل','نسبة الخصم','المحصل لحساب الضريبة']
    alias={'م':['م','ت','مسلسل','الرقم','رقم'],'رقم التسجيل الضريبي':['رقم التسجيل الضريبي','الرقم الضريبي','التسجيل','رقم تسجيل','tax_no','tax_number'],'اسم الممول':['اسم الممول','اسم المورد','الممول','المورد','الاسم','اسم'],'تاريخ التعامل':['تاريخ التعامل','التاريخ','تاريخ','date'],'طبيعة التعامل':['طبيعة التعامل','النوع','طبيعة','نوع التعامل'],'القيمة الإجمالية للتعامل':['القيمة الإجمالية للتعامل','القيمة الإجمالية','قيمة التعامل','القيمة','value'],'نسبة الخصم':['نسبة الخصم','الخصم','نسبة','discount'],'المحصل لحساب الضريبة':['المحصل لحساب الضريبة','المحصل لساب الضريبة','المحصل','ضريبة','tax']}
    hdr=[str(v).strip() for v in df.iloc[0].tolist()]
    mapped={}
    used_cols=set()
    for tgt,al in alias.items():
        for i,h in enumerate(hdr):
            if i in used_cols: continue
            if any(a in h for a in al if a):
                mapped[tgt]=i;used_cols.add(i);break
    if len(mapped)>=6:
        df=df.iloc[1:].reset_index(drop=True)
        for tgt in target:
            if tgt not in mapped:
                mapped[tgt]=len(df.columns);df[len(df.columns)]=''
        df=df[[mapped[t] for t in target]]
        df.columns=target
    else:
        df=df.iloc[:,:len(target)]
        if len(df.columns)<len(target):
            for i in range(len(df.columns),len(target)): df[i]=''
        df.columns=target[:len(df.columns)]
    return df
def read_vat_excel(f):
    df=pd.read_excel(f,header=None,engine='openpyxl')
    cols=['م','اسم الممول','رقم التسجيل الضريبي','ضريبة الجدول','20% قيمة مضافة']
    df=df.iloc[:,:len(cols)]
    if len(df.columns)<len(cols):
        for i in range(len(df.columns),len(cols)): df[i]=''
    df.columns=cols[:len(df.columns)]
    if len(df)>0 and str(df.iloc[0].get('رقم التسجيل الضريبي','')).strip() in ('رقم التسجيل الضريبي','اسم المورد',''):
        df=df.iloc[1:].reset_index(drop=True)
    return df
def safe_val(r,i,d=0):
    if i<len(r) and r[i] is not None:
        try: return float(r[i])
        except: return d
    return d
def safe_str(r,i,d=''):
    if i<len(r) and r[i] is not None: return str(r[i])
    return d
def read_detailed_receipt(f):
    wb=openpyxl.load_workbook(f,data_only=True);ws=wb.active;rows=[]
    for r in ws.iter_rows(min_row=1,values_only=True):
        ic=safe_str(r,19,'').strip()
        if not ic: continue
        name=''
        for i in range(9,min(19,len(r))):
            if r[i] is not None: name=str(r[i]);break
        price=0
        for i in [7,8]:
            if i<len(r) and r[i] is not None:
                try: price=float(r[i])
                except: price=0
                break
        qty=safe_val(r,6)
        disc=0
        for i in [4,5]:
            if i<len(r) and r[i] is not None:
                try: disc=float(r[i])
                except: disc=0
                break
        tax=0
        for i in [1,2]:
            if i<len(r) and r[i] is not None:
                try: tax=float(r[i])
                except: tax=0
                break
        rows.append({'internal_code':ic,'item_name':name,'price':price,'quantity':qty,'discount':disc,'tax':tax})
    wb.close()
    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=['internal_code','item_name','price','quantity','discount','tax'])
def read_barcodes(f):
    wb=openpyxl.load_workbook(f,data_only=True);ws=wb.active;bm={}
    for r in ws.iter_rows(min_row=1,values_only=True):
        c,b=safe_str(r,1,'').strip(),safe_str(r,6,'').strip()
        if c and b: bm[c]=b
    wb.close()
    return bm
def gen_template(ddf,bmap):
    tp=os.path.join(DATA_DIR,'Template_Portal.xlsx')
    if os.path.exists(tp): wb=openpyxl.load_workbook(tp)
    else:
        wb=openpyxl.Workbook();ws=wb.active;ws.title="بنود الفاتورة"
        for c,h in enumerate(['كود الصنف','الكود الداخلي','الوصف','كود الوحدة','السعر','الكمية','الخصم','خصم الأصناف','كود ضريبة 1','نسبة الضريبة 1','كود ضريبة 2','نسبة الضريبة 2','كود ضريبة 3','نسبة الضريبة 3','كود ضريبة 4','نسبة الضريبة 4','كود ضريبة 5','نسبة الضريبة 5','كود ضريبة 6','نسبة الضريبة 6','العملة','سعر العملة','عرض الأكواد'],1):
            ws.cell(row=1,column=c,value=h)
        wb.save(tp)
    ws=wb.active
    for ri in range(ws.max_row,1,-1): ws.delete_rows(ri)
    for idx,(_,r) in enumerate(ddf.iterrows()):
        i=idx+2;ic=str(r['internal_code']).strip();bc=bmap.get(ic,'')
        p=float(r['price']) if r['price'] else 0;q=float(r['quantity']) if r['quantity'] else 0
        d=float(r['discount']) if r['discount'] else 0;t=float(r['tax']) if r['tax'] else 0
        tc,tr='',''
        if t>0: p/=1.14;tc='V009';tr=14;
        if d>0 and t>0: d/=1.14
        ws.cell(row=i,column=1,value=bc);ws.cell(row=i,column=2,value=ic);ws.cell(row=i,column=3,value=r['item_name'])
        ws.cell(row=i,column=4,value='EA');ws.cell(row=i,column=5,value=math.ceil(p*100000)/100000 if p>0 else 0);ws.cell(row=i,column=6,value=q)
        ws.cell(row=i,column=7,value=math.ceil(d*100000)/100000 if d>0 else '');ws.cell(row=i,column=8,value='')
        ws.cell(row=i,column=9,value=tc);ws.cell(row=i,column=10,value=tr)
    out=BytesIO();wb.save(out);out.seek(0);wb.close()
    return out
def fmt(n):
    try: return f"{float(n):,.2f}"
    except: return "0.00"
MONTHS={1:'يناير',2:'فبراير',3:'مارس',4:'أبريل',5:'مايو',6:'يونيو',7:'يوليو',8:'أغسطس',9:'سبتمبر',10:'أكتوبر',11:'نوفمبر',12:'ديسمبر'}
REQUESTS_FILE=os.path.join(DATA_DIR,"requests_registry.json")

def load_requests():
    gh=gh_read("requests_registry.json")
    if gh is not None:
        with open(REQUESTS_FILE,'w',encoding='utf-8') as f: json.dump(gh,f,ensure_ascii=False,indent=2,default=str)
        return gh
    if os.path.exists(REQUESTS_FILE):
        try:
            with open(REQUESTS_FILE,'r',encoding='utf-8') as f: return json.load(f)
        except: return []
    return []

def save_requests(data):
    with open(REQUESTS_FILE,'w',encoding='utf-8') as f: json.dump(data,f,ensure_ascii=False,indent=2,default=str)
    gh_write("requests_registry.json",data)


def _replace_in_paragraph(para, old, new):
    full = para.text
    if old in full:
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return True
        combined = ''.join(r.text for r in para.runs)
        if old in combined:
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = combined.replace(old, new)
                else:
                    run.text = ''
            return True
    return False

def _set_para_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
    else:
        para.clear()
        para.add_run(new_text)

def gen_form41_word(supplier_name, tax_number, results, export_date_str, request_date_str=None, request_number=None):
    template_path=os.path.join(DATA_DIR,"جواب نموذج 41.docx")
    if not os.path.exists(template_path):
        return None, "ملف القالب غير موجود"
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
        doc=Document(template_path)
        req=str(request_number) if request_number else tax_number
        formatted_tax=tax_number
        if len(tax_number)==9:
            formatted_tax=f"{tax_number[6:]}-{tax_number[3:6]}-{tax_number[:3]}"
        for para in doc.paragraphs:
            txt=para.text
            if 'القيد' in txt and 'مالي' in txt:
                _set_para_text(para, f'القيد : {req} / مالي / 2026')
            elif 'التاريخ' in txt and 'التاريخ :' in txt:
                _set_para_text(para, f'التاريخ : {export_date_str}')
            elif 'إلى /' in txt and 'رقم تسجيل ضريبي' in txt:
                new_text=f'إلى / {supplier_name} برقم تسجيل ضريبي ({formatted_tax})'
                _set_para_text(para, new_text)
                pf=para.paragraph_format
                pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE
                pf.line_spacing=1.5
                pf.space_before=Pt(0)
                pf.space_after=Pt(0)
        if request_date_str:
            for para in doc.paragraphs:
                _replace_in_paragraph(para, 'بتاريخ 8/4/2026', f'بتاريخ {request_date_str}')
        t=doc.tables[0]
        existing_rows=len(t.rows)-1
        needed=len(results)
        if needed>existing_rows:
            for _ in range(needed-existing_rows):
                src_row=t.rows[-1]
                new_row=t.add_row()
                for ci,cell in enumerate(new_row.cells):
                    cell.text=''
        elif needed<existing_rows:
            for _ in range(existing_rows-needed):
                tr=t.rows[-1]
                tbl=t._tbl
                tbl.remove(tr._tr)
        for idx,r in enumerate(results):
            ri=idx+1
            period_val=str(r.get('الفترة','7/2025'))
            period_parts=period_val.split('/')
            month_str=MONTHS.get(int(period_parts[0]),'')
            year_str=period_parts[1] if len(period_parts)>1 else '2025'
            period=f"{month_str} {year_str}"
            tax_val=_sf(r.get('المحصل لحساب الضريبة',0))
            pay_num=str(r.get('رقم المدفوعة',''))
            pay_date=str(r.get('تاريخ المدفوعة',''))[:10]
            if pay_date and '-' in pay_date:
                parts=pay_date.split('-')
                if len(parts)==3: pay_date=f"{parts[2]}/{parts[1]}/{parts[0]}"
            elif pay_date and '/' in pay_date:
                parts=pay_date.split('/')
                if len(parts)==3 and len(parts[0])==4: pay_date=f"{parts[2]}/{parts[1]}/{parts[0]}"
            t.cell(ri,0).text=str(idx+1)
            t.cell(ri,1).text=str(int(tax_val)) if tax_val==int(tax_val) else str(tax_val)
            t.cell(ri,2).text=period
            t.cell(ri,3).text=pay_num
            t.cell(ri,4).text=pay_date
        out_path=os.path.join(DATA_DIR,f"جواب_نموذج_41_{tax_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(out_path)
        return out_path, None
    except Exception as e:
        return None, str(e)

def gen_vat_word(supplier_name, tax_number, results, export_date_str, request_date_str=None, request_number=None):
    template_path=os.path.join(DATA_DIR,"جواب نموذج 41.docx")
    if not os.path.exists(template_path):
        return None, "ملف القالب غير موجود"
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
        doc=Document(template_path)
        req=str(request_number) if request_number else tax_number
        formatted_tax=tax_number
        if len(tax_number)==9:
            formatted_tax=f"{tax_number[6:]}-{tax_number[3:6]}-{tax_number[:3]}"
        for para in doc.paragraphs:
            txt=para.text
            if 'القيد' in txt and 'مالي' in txt:
                _set_para_text(para, f'القيد : {req} / مالي / 2026')
            elif 'التاريخ' in txt and 'التاريخ :' in txt:
                _set_para_text(para, f'التاريخ : {export_date_str}')
            elif 'إلى /' in txt and 'رقم تسجيل ضريبي' in txt:
                new_text=f'إلى / {supplier_name} برقم تسجيل ضريبي ({formatted_tax})'
                _set_para_text(para, new_text)
                pf=para.paragraph_format
                pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE
                pf.line_spacing=1.5
                pf.space_before=Pt(0)
                pf.space_after=Pt(0)
        if request_date_str:
            for para in doc.paragraphs:
                _replace_in_paragraph(para, 'بتاريخ 8/4/2026', f'بتاريخ {request_date_str}')
        t=doc.tables[0]
        hdr_cell=t.cell(0,2)
        for p in hdr_cell.paragraphs:
            for run in p.runs:
                run.text=run.text.replace('نموذج 41 شهر','النوع')
            if 'نموذج 41 شهر' in p.text:
                p.clear()
                p.add_run('النوع')
        existing_rows=len(t.rows)-1
        needed=len(results)
        if needed>existing_rows:
            for _ in range(needed-existing_rows):
                new_row=t.add_row()
                for ci,cell in enumerate(new_row.cells):
                    cell.text=''
        elif needed<existing_rows:
            for _ in range(existing_rows-needed):
                tr=t.rows[-1]
                tbl=t._tbl
                tbl.remove(tr._tr)
        for idx,r in enumerate(results):
            ri=idx+1
            tv_val=_sf(r.get('20% قيمة مضافة',0))
            tt_val=_sf(r.get('ضريبة الجدول',0))
            if tv_val>0:
                ntype='20% قيمة مضافة'
                tax_val=tv_val
            else:
                ntype='جدول'
                tax_val=tt_val
            pay_num=str(r.get('رقم المدفوعة',''))
            pay_date=str(r.get('تاريخ المدفوعة',''))[:10]
            if pay_date and '-' in pay_date:
                parts=pay_date.split('-')
                if len(parts)==3: pay_date=f"{parts[2]}/{parts[1]}/{parts[0]}"
            elif pay_date and '/' in pay_date:
                parts=pay_date.split('/')
                if len(parts)==3 and len(parts[0])==4: pay_date=f"{parts[2]}/{parts[1]}/{parts[0]}"
            t.cell(ri,0).text=str(idx+1)
            t.cell(ri,1).text=str(int(tax_val)) if tax_val==int(tax_val) else str(tax_val)
            t.cell(ri,2).text=ntype
            t.cell(ri,3).text=pay_num
            t.cell(ri,4).text=pay_date
        out_path=os.path.join(DATA_DIR,f"جواب_قيمة_مضافة_{tax_number}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(out_path)
        return out_path, None
    except Exception as e:
        return None, str(e)
        if os.path.exists(tmp_path):
            try: os.remove(tmp_path)
            except: pass

# ====================== HELPER: BATCH CARD HTML ======================
def batch_card_html(rec, btype):
    icon = "📋" if btype=='f41' else "💰"
    cls = "b-f41" if btype=='f41' else "b-vat"
    color = "#6c5ce7" if btype=='f41' else "#00cec9"
    label = "نموذج 41" if btype=='f41' else "القيمة المضافة"
    cnt = len(rec.get('records',[]))
    return f"""<div class="erp-batch {cls}">
    <div class="erp-batch-head">
        <div class="erp-batch-icon" style="background:{color}15;">{icon}</div>
        <div><div class="erp-batch-title">{rec.get('file_name','')}</div><div class="erp-batch-sub">{label}</div></div>
    </div>
    <div class="erp-batch-grid">
        <div><div class="erp-batch-k">الفترة</div><div class="erp-batch-v">{rec['model_month']}/{rec['model_year']}</div></div>
        <div><div class="erp-batch-k">رقم المدفوعة</div><div class="erp-batch-v">{rec.get('payment_number','-')}</div></div>
        <div><div class="erp-batch-k">تاريخ المدفوعة</div><div class="erp-batch-v">{_fmt_date_dmy(rec.get('payment_date','-'))}</div></div>
        <div><div class="erp-batch-k">تاريخ الرفع</div><div class="erp-batch-v">{_fmt_date_dmy(rec.get('upload_date',''))}</div></div>
    </div>
    <div class="erp-batch-foot"><span class="erp-batch-count"><span>{cnt}</span> سجل</span></div>
</div>"""

def _fmt_date_dmy(val):
    s=str(val)[:10] if val else '-'
    if len(s)==10 and s[4]=='-' and s[7]=='-':
        return f"{s[8:10]}/{s[5:7]}/{s[:4]}"
    return s

def _sf(v):
    try:
        import math
        f=float(v) if v is not None and v!='' else 0.0
        return 0.0 if math.isnan(f) or math.isinf(f) else f
    except: return 0.0
def meta_html(label, value, color="var(--accent2)"):
    return f"""<div class="erp-meta-item"><div class="erp-meta-k">{label}</div><div class="erp-meta-v" style="color:{color}">{value}</div></div>"""

# ====================== ADMIN ======================
if page == ADMIN_PAGE:
    cu=get_current_user()
    if cu.get('role')!='admin':
        st.error("لا تملك صلاحية الوصول");st.stop()
    st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>إدارة المستخدمين</h3></div>',unsafe_allow_html=True)
    users=load_users()
    tab_add,tab_list,tab_edit=st.tabs(["➕ إضافة مستخدم","📋 قائمة المستخدمين","✏️ تعديل الصلاحيات"])
    with tab_add:
        with st.form("add_user_form",clear_on_submit=True):
            c1,c2=st.columns(2)
            with c1: new_user=st.text_input("اسم المستخدم (Username)")
            with c2: new_name=st.text_input("الاسم الكامل")
            c3,c4=st.columns(2)
            with c3: new_pw=st.text_input("كلمة المرور",type="password")
            with c4: new_role=st.selectbox("الدور",["user","admin"])
            st.markdown('<p style="color:var(--text2);font-size:.8rem;margin:.5rem 0 .3rem;">الصلاحيات:</p>',unsafe_allow_html=True)
            perm_cols=st.columns(4)
            perm_checks=[]
            for i,pg in enumerate(ALL_PAGES):
                with perm_cols[i%4]:
                    perm_checks.append((pg,st.checkbox(pg,key=f"perm_{i}")))
            if st.form_submit_button("💾 إضافة المستخدم",type="primary",use_container_width=True):
                if not new_user or not new_pw or not new_name:
                    st.error("أدخل كل البيانات المطلوبة")
                elif any(u['username']==new_user for u in users):
                    st.error("اسم المستخدم موجود بالفعل")
                else:
                    perms=[p for p,c in perm_checks if c]
                    if new_role=='admin': perms=ALL_PAGES+[ADMIN_PAGE]
                    users.append({"username":new_user.strip(),"password":_hash_pw(new_pw),"display_name":new_name.strip(),"role":new_role,"permissions":perms,"created_at":datetime.now().isoformat()})
                    save_users(users)
                    st.success(f"تم إضافة {new_name.strip()} بنجاح!")
    with tab_list:
        if users:
            for idx,u in enumerate(users):
                is_admin=u.get('role')=='admin'
                role_label="👑 مدير" if is_admin else "👤 مستخدم"
                perms=u.get('permissions',[])
                perms_display="، ".join(perms) if perms else "بدون صلاحيات"
                with st.container():
                    st.markdown(f"""<div style="background:var(--surface);border:1px solid var(--border);border-radius:14px;padding:1rem 1.2rem;margin-bottom:.6rem;display:flex;align-items:center;justify-content:space-between;">
                        <div><strong style="color:var(--text);font-size:.9rem;">{u.get('display_name','')}</strong> <span style="color:rgba(255,255,255,.3);font-size:.75rem;">@{u['username']}</span></div>
                        <span style="color:var(--accent2);font-size:.75rem;">{role_label}</span></div>""",unsafe_allow_html=True)
                    st.markdown(f'<p style="color:var(--text2);font-size:.7rem;margin:-.5rem 0 .5rem;">الصلاحيات: {perms_display}</p>',unsafe_allow_html=True)
                    ec1,ec2=st.columns([3,1])
                    with ec2:
                        if u['username']!='admin' and st.button("🗑️ حذف",key=f"del_user_{idx}",type="secondary",use_container_width=True):
                            users=[x for x in users if x['username']!=u['username']]
                            save_users(users);st.success(f"تم حذف {u.get('display_name','')}");st.rerun()
        else:
            st.info("لا يوجد مستخدمون")
    with tab_edit:
        editable_users=[u for u in users if u.get('role')!='admin']
        if not editable_users:
            st.info("لا يوجد مستخدمون للتعديل")
        else:
            edit_user_names=[f"{u.get('display_name','')} (@{u['username']})" for u in editable_users]
            sel_idx=st.selectbox("اختر المستخدم",range(len(edit_user_names)),format_func=lambda i:edit_user_names[i],key="edit_perm_user")
            sel_user=editable_users[sel_idx]
            st.markdown(f'<p style="color:var(--text);font-size:.85rem;margin:.5rem 0;">صلاحيات <strong style="color:var(--accent2);">{sel_user.get("display_name","")}</strong></p>',unsafe_allow_html=True)
            cur_perms=sel_user.get('permissions',[])
            new_perms=[]
            perm_cols=st.columns(3)
            for i,pg in enumerate(ALL_PAGES):
                with perm_cols[i%3]:
                    new_perms.append((pg,st.checkbox(pg,value=pg in cur_perms,key=f"edit_perm_{sel_user['username']}_{i}")))
            if st.button("💾 حفظ الصلاحيات",key="save_edit_perms",type="primary",use_container_width=True):
                final_perms=[p for p,c in new_perms if c]
                for u in users:
                    if u['username']==sel_user['username']:
                        u['permissions']=final_perms
                save_users(users)
                cu=get_current_user()
                if cu and cu.get('username')==sel_user['username']:
                    cu['permissions']=final_perms
                st.success(f"تم تعديل صلاحيات {sel_user.get('display_name','')} بنجاح!")
                st.rerun()
    st.stop()

# ====================== HOME ======================
if page == "🏠 الرئيسية":
    # Detail view
    if 'detail_view' in st.session_state:
        dv = st.session_state['detail_view']
        f = FORM41_FILE if dv['type']=='f41' else VAT_FILE
        data = load_data(f)
        if dv['index'] >= len(data):
            st.error("السجل غير موجود"); del st.session_state['detail_view']; st.rerun()
        rec = data[dv['index']]
        records = rec.get('records',[])
        tlabel = "نموذج 41" if dv['type']=='f41' else "القيمة المضافة"

        st.markdown(f"""<div class="erp-detail-header">
            <a class="erp-back" onclick="window.parent.document.querySelector('[data-testid=stSidebar]').style.display='block'">← رجوع</a>
            <span class="erp-detail-title">{tlabel} — {rec.get('file_name','')}</span>
        </div>""", unsafe_allow_html=True)

        if st.button("← رجوع للرئيسية", key="back_dash"):
            del st.session_state['detail_view']; st.rerun()

        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>بيانات الربط</h3></div>', unsafe_allow_html=True)
        st.markdown(f"""<div class="erp-meta-grid">
            {meta_html("الملف", rec.get('file_name',''))}
            {meta_html("الفترة", f"{rec['model_month']}/{rec['model_year']}", "#55efc4")}
            {meta_html("رقم المدفوعة", rec.get('payment_number','-'), "#fdcb6e")}
            {meta_html("تاريخ المدفوعة", _fmt_date_dmy(rec.get('payment_date','-')), "#74b9ff")}
            {meta_html("عدد السجلات", str(len(records)), "#a29bfe")}
            {meta_html("تاريخ الرفع", _fmt_date_dmy(rec.get('upload_date','')))}
        </div>""", unsafe_allow_html=True)

        cu=get_current_user()
        if cu and cu.get('role')=='admin':
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>تعديل البيانات (Admin)</h3></div>', unsafe_allow_html=True)
            ec1,ec2=st.columns(2)
            with ec1: new_pn=st.text_input("رقم المدفوعة",value=rec.get('payment_number',''),key="edit_pn")
            with ec2:
                try: pd_val=datetime.fromisoformat(str(rec.get('payment_date',''))[:10]).date() if rec.get('payment_date') else datetime.now().date()
                except: pd_val=datetime.now().date()
                new_pd=st.date_input("تاريخ المدفوعة",value=pd_val,key="edit_pd")
            em1,ey1=st.columns(2)
            with em1: new_mm=st.selectbox("شهر النموذج",range(1,13),index=rec.get('model_month',7)-1,format_func=lambda x:f"{x} - {MONTHS[x]}",key="edit_mm")
            with ey1: new_yy=st.selectbox("سنة النموذج",range(2020,2031),index=rec.get('model_year',2025)-2020,key="edit_yy")
            e1,e2=st.columns(2)
            with e1:
                if st.button("💾 حفظ التعديل",key="save_edit",type="primary",use_container_width=True):
                    data[dv['index']]['payment_number']=new_pn
                    data[dv['index']]['payment_date']=new_pd.isoformat()
                    data[dv['index']]['model_month']=new_mm
                    data[dv['index']]['model_year']=new_yy
                    save_data(f,data)
                    st.success("تم التعديل بنجاح!");st.rerun()
            with e2:
                if st.button("🗑️ حذف النموذج بالكامل",key="delete_batch",type="secondary",use_container_width=True):
                    st.session_state['confirm_delete_batch']=dv['index']
                    st.session_state['confirm_delete_file']=f
            if st.session_state.get('confirm_delete_batch')==dv['index'] and st.session_state.get('confirm_delete_file')==f:
                st.markdown('<div style="padding:.8rem 1rem;border-radius:10px;background:rgba(255,107,107,.08);border:1px solid rgba(255,107,107,.2);margin:.5rem 0;">⚠ هل أنت متأكد من حذف هذا النموذج بالكامل؟ لا يمكن التراجع.</div>', unsafe_allow_html=True)
                cd1,cd2=st.columns(2)
                with cd1:
                    if st.button("نعم، احذف",key="confirm_del_yes",type="primary",use_container_width=True):
                        del_data=load_data(f)
                        del_data.pop(dv['index'])
                        save_data(f,del_data)
                        del st.session_state['confirm_delete_batch']
                        del st.session_state['confirm_delete_file']
                        del st.session_state['detail_view']
                        st.success("تم الحذف!");st.rerun()
                with cd2:
                    if st.button("إلغاء",key="confirm_del_no",use_container_width=True):
                        del st.session_state['confirm_delete_batch']
                        del st.session_state['confirm_delete_file']
                        st.rerun()

        if dv['type'] == 'f41':
            total_tax = sum(_sf(r.get('المحصل لحساب الضريبة',0)) for r in records)
            total_val = sum(_sf(r.get('القيمة الإجمالية للتعامل',0)) for r in records)
            s1,s2,s3 = st.columns(3)
            with s1: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">إجمالي المحصل لحساب الضريبة</div><div class="erp-stat-value">{fmt(total_tax)}</div></div>', unsafe_allow_html=True)
            with s2: st.markdown(f'<div class="erp-stat s-cyan"><div class="erp-stat-label">إجمالي القيمة الإجمالية للتعامل</div><div class="erp-stat-value">{fmt(total_val)}</div></div>', unsafe_allow_html=True)
            with s3: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">عدد السجلات</div><div class="erp-stat-value">{len(records)}</div></div>', unsafe_allow_html=True)
        else:
            tt=sum(_sf(r.get('ضريبة الجدول',0)) for r in records)
            tv=sum(_sf(r.get('20% قيمة مضافة',0)) for r in records)
            s1,s2,s3=st.columns(3)
            with s1: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">ضريبة الجدول</div><div class="erp-stat-value">{fmt(tt)}</div></div>', unsafe_allow_html=True)
            with s2: st.markdown(f'<div class="erp-stat s-pink"><div class="erp-stat-label">20% قيمة مضافة</div><div class="erp-stat-value">{fmt(tv)}</div></div>', unsafe_allow_html=True)
            with s3: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الإجمالي</div><div class="erp-stat-value">{fmt(tt+tv)}</div></div>', unsafe_allow_html=True)

        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>البيانات التفصيلية</h3></div>', unsafe_allow_html=True)
        st.markdown('<div class="erp-card">', unsafe_allow_html=True)
        st.dataframe(pd.DataFrame(records), use_container_width=True, height=400)
        st.markdown('</div>', unsafe_allow_html=True)
        st.stop()

    # Dashboard
    f41 = load_data(FORM41_FILE); vat = load_data(VAT_FILE)
    f41_n = sum(len(r.get('records',[])) for r in f41)
    vat_n = sum(len(r.get('records',[])) for r in vat)
    portal_out = load_data(PORTAL_OUT_FILE); portal_in = load_data(PORTAL_IN_FILE)
    portal_out_n = sum(r.get('records_count',len(r.get('records',[]))) for r in portal_out)
    portal_in_n = sum(r.get('records_count',len(r.get('records',[]))) for r in portal_in)

    st.markdown(f"""<div class="erp-topbar"><div><h2>{page}</h2><p>مرحباً بك في لوحة التحكم</p></div>
<div class="erp-topbar-right"><span class="erp-badge">📊 Dashboard</span><a href="https://invoicing.eta.gov.eg/" target="_blank" style="background:rgba(0,206,201,.12);border:1px solid rgba(0,206,201,.3);border-radius:10px;padding:.35rem .9rem;color:#00cec9;font-size:.72rem;font-weight:600;text-decoration:none;cursor:pointer;transition:all .3s;">portal الفواتير الإلكترونية</a></div></div>""", unsafe_allow_html=True)

    st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>نظرة عامة</h3></div>', unsafe_allow_html=True)
    c1,c2,c3,c4,c5=st.columns(5)
    with c1: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">نموذج 41</div><div class="erp-stat-value">{f41_n}</div><div class="erp-stat-sub">{len(f41)} رفع</div></div>', unsafe_allow_html=True)
    with c2: st.markdown(f'<div class="erp-stat s-cyan"><div class="erp-stat-label">القيمة المضافة</div><div class="erp-stat-value">{vat_n}</div><div class="erp-stat-sub">{len(vat)} رفع</div></div>', unsafe_allow_html=True)
    with c3: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">فواتير صادرة</div><div class="erp-stat-value">{portal_out_n}</div><div class="erp-stat-sub">{len(portal_out)} رفع</div></div>', unsafe_allow_html=True)
    with c4: st.markdown(f'<div class="erp-stat s-pink"><div class="erp-stat-label">فواتير واردة</div><div class="erp-stat-value">{portal_in_n}</div><div class="erp-stat-sub">{len(portal_in)} رفع</div></div>', unsafe_allow_html=True)
    with c5: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الإجمالي</div><div class="erp-stat-value">{f41_n+vat_n+portal_out_n+portal_in_n}</div><div class="erp-stat-sub">سجل</div></div>', unsafe_allow_html=True)

    # F41 batches
    if f41:
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>نماذج 41 المرفوعة</h3></div>', unsafe_allow_html=True)
        cols_per_row = min(len(f41), 3)
        for start in range(0, len(f41), 3):
            cols = st.columns(3)
            for ci, idx in enumerate(range(start, min(start+3, len(f41)))):
                with cols[ci]:
                    rec = f41[idx]
                    st.markdown(batch_card_html(rec, 'f41'), unsafe_allow_html=True)
                    if st.button("📊 عرض التفاصيل", key=f"view_f41_{idx}", type="primary", use_container_width=True):
                        st.session_state['detail_view'] = {'type':'f41','index':idx}
                        st.rerun()

    if vat:
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>نماذج القيمة المضافة المرفوعة</h3></div>', unsafe_allow_html=True)
        for start in range(0, len(vat), 3):
            cols = st.columns(3)
            for ci, idx in enumerate(range(start, min(start+3, len(vat)))):
                with cols[ci]:
                    rec = vat[idx]
                    st.markdown(batch_card_html(rec, 'vat'), unsafe_allow_html=True)
                    if st.button("📊 عرض التفاصيل", key=f"view_vat_{idx}", type="primary", use_container_width=True):
                        st.session_state['detail_view'] = {'type':'vat','index':idx}
                        st.rerun()

    if not f41 and not vat:
        st.markdown('<div class="erp-empty"><div class="erp-empty-icon">📋</div><h3>لا توجد بيانات بعد</h3><p>ابدأ برفع الملفات من القائمة الجانبية</p></div>', unsafe_allow_html=True)

# ====================== FORM 41 ======================
elif page == "📋 نموذج 41":
    if not user_has_permission(page): st.error("لا تملك صلاحية الوصول");st.stop()
    sub = st.radio("f41", ["📤 رفع وربط البيانات","🔍 استعلام بالسجل الضريبي","📅 استعلام بالفترة"], horizontal=True, label_visibility="collapsed")

    if sub == "📤 رفع وربط البيانات":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>رفع شيت نموذج 41</h3></div>', unsafe_allow_html=True)
        st.markdown('<div class="erp-info"><strong>الأعمدة:</strong> م • رقم التسجيل الضريبي • اسم الممول • تاريخ التعامل • طبيعة التعامل • القيمة الإجمالية للتعامل • نسبة الخصم • المحصل لحساب الضريبة</div>', unsafe_allow_html=True)
        st.markdown('<div class="erp-card">', unsafe_allow_html=True)
        up=st.file_uploader("ارفع شيت نموذج 41",type=['xlsx','xls'],key="f41_up",label_visibility="collapsed")
        if up:
            if 'f41_df' not in st.session_state or st.session_state.get('f41_file')!=up.name:
                st.session_state['f41_df']=read_form41_excel(up);st.session_state['f41_file']=up.name
            df=st.session_state['f41_df']
            st.markdown(f'<div style="margin:.5rem 0;padding:.6rem 1rem;border-radius:10px;background:rgba(0,184,148,.08);border:1px solid rgba(0,184,148,.15);color:#55efc4;font-size:.82rem;">✓ تم الرفع — {len(df)} سجل</div>', unsafe_allow_html=True)
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>حذف سطور (اختياري)</h3></div>', unsafe_allow_html=True)
            row_labels=[f"سطر {i+1} — {df.iloc[i].to_dict().get('اسم الممول','')}" for i in range(len(df))]
            rm=st.multiselect("اختر السطور اللي عايز تشيلها",options=list(range(len(df))),format_func=lambda i:row_labels[i],key="f41_rm")
            if rm:
                st.markdown(f'<div style="padding:.5rem 1rem;border-radius:8px;background:rgba(255,107,107,.08);border:1px solid rgba(255,107,107,.15);color:#ff6b6b;font-size:.8rem;margin:.5rem 0;">⚠ سيتم حذف {len(rm)} سطر — المتبقي: {len(df)-len(rm)}</div>', unsafe_allow_html=True)
                if st.button("🗑️ تطبيق الحذف",key="f41_arm"):
                    st.session_state['f41_df']=df.drop(rm).reset_index(drop=True);st.rerun()
            df=st.session_state['f41_df']
            st.dataframe(df,use_container_width=True,height=280)
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>ربط البيانات</h3></div>',unsafe_allow_html=True)
            c1,c2=st.columns(2)
            with c1: f41_pd=st.date_input("تاريخ المدفوعة",key="f41_pd")
            with c2: f41_pn=st.text_input("رقم المدفوعة",key="f41_pn")
            c3,c4=st.columns(2)
            with c3: f41_m=st.selectbox("شهر النموذج",range(1,13),format_func=lambda x:f"{x} - {MONTHS[x]}",key="f41_m")
            with c4: f41_y=st.selectbox("سنة النموذج",range(2020,2031),key="f41_y")
            if st.button("💾 حفظ وربط البيانات",key="f41_save",type="primary"):
                data=load_data(FORM41_FILE)
                data.append({"id":str(uuid.uuid4()),"upload_date":datetime.now().isoformat(),"payment_date":f41_pd.isoformat(),"payment_number":f41_pn,"model_month":f41_m,"model_year":f41_y,"file_name":up.name,"records":df.to_dict('records')})
                save_data(FORM41_FILE,data);st.success("تم الحفظ بنجاح!");del st.session_state['f41_df']
        st.markdown('</div>',unsafe_allow_html=True)

    elif sub == "🔍 استعلام بالسجل الضريبي":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>استعلام بالسجل الضريبي</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        c1,c2,c3=st.columns(3)
        with c1: qt=st.text_input("رقم التسجيل الضريبي",key="f41_qt")
        with c2: qm=st.selectbox("الشهر",range(1,13),index=6,format_func=lambda x:f"{x}-{MONTHS[x]}",key="f41_qm")
        with c3: qy=st.selectbox("السنة",range(2020,2031),key="f41_qy")
        if st.button("🔍 بحث",key="f41_bt",type="primary"):
            if not qt.strip(): st.warning("أدخل رقم التسجيل الضريبي")
            else:
                data=load_data(FORM41_FILE);res=[]
                for rec in data:
                    if rec['model_month']==qm and rec['model_year']==qy:
                        for row in rec.get('records',[]):
                            if str(row.get('رقم التسجيل الضريبي','')).strip()==qt.strip():
                                rc=dict(row);rc['فترة الخصم']=f"{qm}/{qy}";rc['رقم المدفوعة']=rec.get('payment_number','');rc['تاريخ المدفوعة']=rec.get('payment_date','');res.append(rc)
                if res:
                    sn=res[0].get('اسم الممول','');td=sum(_sf(r.get('المحصل لحساب الضريبة',0)) for r in res)
                    s1,s2,s3,s4=st.columns(4)
                    with s1: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">اسم الممول</div><div class="erp-stat-value" style="font-size:.95rem">{sn}</div></div>',unsafe_allow_html=True)
                    with s2: st.markdown(f'<div class="erp-stat s-cyan"><div class="erp-stat-label">السجل الضريبي</div><div class="erp-stat-value" style="font-size:.95rem">{qt}</div></div>',unsafe_allow_html=True)
                    with s3: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">الخصومات</div><div class="erp-stat-value">{fmt(td)}</div></div>',unsafe_allow_html=True)
                    with s4: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الفترة</div><div class="erp-stat-value" style="font-size:1.1rem">{qm}/{qy}</div></div>',unsafe_allow_html=True)
                    st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>التفاصيل</h3></div>',unsafe_allow_html=True)
                    st.dataframe(pd.DataFrame(res),use_container_width=True)
                else: st.info("لم يتم العثور على نتائج")
        st.markdown('</div>',unsafe_allow_html=True)

    elif sub == "📅 استعلام بالفترة":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>استعلام بالفترة</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        qt=st.text_input("رقم التسجيل الضريبي",key="f41_pqt")
        periods_list=[(m,y) for y in range(2020,2031) for m in range(1,13)]
        period_labels={f"{m}/{y}":f"{m}-{MONTHS[m]} {y}" for m,y in periods_list}
        sel=st.multiselect("اختر الفترات",options=list(period_labels.keys()),format_func=lambda x:period_labels[x],default=[],key="f41_periods")
        if st.button("🔍 بحث",key="f41_bp",type="primary"):
            if not sel: st.warning("اختر فترة واحدة على الأقل")
            else:
                periods=set()
                for s in sel:
                    parts=s.split('/')
                    periods.add((int(parts[0]),int(parts[1])))
                data=load_data(FORM41_FILE);res=[]
                for rec in data:
                    if (rec['model_month'],rec['model_year']) in periods:
                        for row in rec.get('records',[]):
                            if qt.strip():
                                if str(row.get('رقم التسجيل الضريبي','')).strip()==qt.strip():
                                    rc=dict(row);rc['الفترة']=f"{rec['model_month']}/{rec['model_year']}";rc['رقم المدفوعة']=rec.get('payment_number','');rc['تاريخ المدفوعة']=rec.get('payment_date','');res.append(rc)
                            else:
                                rc=dict(row);rc['الفترة']=f"{rec['model_month']}/{rec['model_year']}";rc['رقم المدفوعة']=rec.get('payment_number','');rc['تاريخ المدفوعة']=rec.get('payment_date','');res.append(rc)
                st.session_state['f41_period_res']=res
                st.session_state['f41_period_sel']=sel
                st.session_state['f41_period_qt']=qt
        res=st.session_state.get('f41_period_res',[])
        sel=st.session_state.get('f41_period_sel',[])
        qt_val=st.session_state.get('f41_period_qt','')
        if res:
            ta=sum(_sf(r.get('المحصل لحساب الضريبة',0)) for r in res)
            st.success(f"{len(res)} سجل")
            s1,s2=st.columns(2)
            with s1: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">العدد</div><div class="erp-stat-value">{len(res)}</div></div>',unsafe_allow_html=True)
            with s2: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الإجمالي</div><div class="erp-stat-value">{fmt(ta)}</div></div>',unsafe_allow_html=True)
            st.dataframe(pd.DataFrame(res),use_container_width=True)
            st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>تصدير جواب نموذج 41</h3></div>',unsafe_allow_html=True)
            export_card='<div class="erp-card"><div class="erp-card-header"><div class="erp-card-icon" style="background:linear-gradient(135deg,rgba(253,203,110,.15),rgba(253,203,110,.03));">📝</div><div><h3>إعدادات التصدير</h3></div></div></div>'
            st.markdown(export_card,unsafe_allow_html=True)
            c1,c2=st.columns(2)
            with c1: export_date=st.date_input("تاريخ تقديم طلب الشهادة",key="f41_ed")
            with c2:
                reqs=load_requests()
                next_num=1
                if reqs:
                    nums=[r.get('request_number',0) for r in reqs]
                    next_num=max(nums)+1 if nums else 1
                req_num=st.number_input("رقم الطلب",min_value=1,value=next_num,step=1,key="f41_rn")
            supplier_name=res[0].get('اسم الممول','')
            tax_number=str(res[0].get('رقم التسجيل الضريبي','')).strip()
            st.markdown(f'<div style="padding:.5rem 1rem;border-radius:10px;background:rgba(108,92,231,.06);border:1px solid rgba(108,92,231,.1);font-size:.82rem;margin:.5rem 0;">المورد: <strong>{supplier_name}</strong> — التسجيل الضريبي: <strong>{tax_number}</strong></div>',unsafe_allow_html=True)
            if st.button("📄 إنشاء جواب نموذج 41",key="f41_export",type="primary"):
                export_date_formatted=export_date.strftime("%d/%m/%Y")
                request_date_formatted=export_date.strftime("%d/%m/%Y")
                with st.spinner("جاري إنشاء الملف..."):
                    out_path,err=gen_form41_word(supplier_name,tax_number,res,export_date_formatted,request_date_formatted,int(req_num))
                if err:
                    st.error(f"خطأ: {err}")
                else:
                    req_record={
                        "id":str(uuid.uuid4()),
                        "request_number":int(req_num),
                        "tax_number":tax_number,
                        "supplier_name":supplier_name,
                        "periods":sorted(list(sel)),
                        "request_date":export_date.isoformat(),
                        "export_date":export_date_formatted,
                        "records_count":len(res),
                        "total_tax":ta,
                        "created_at":datetime.now().isoformat(),
                        "file_name":os.path.basename(out_path)
                    }
                    reqs=load_requests()
                    reqs.append(req_record)
                    save_requests(reqs)
                    st.success(f"تم إنشاء الملف — رقم الطلب: {req_num}")
                    with open(out_path,'rb') as f:
                        st.download_button("📥 تحميل الجواب",data=f.read(),file_name=os.path.basename(out_path),mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
        elif sel:
            st.info("لا توجد سجلات في هذه الفترات")
        st.markdown('</div>',unsafe_allow_html=True)

        st.markdown('<div class="erp-section" style="margin-top:1.5rem"><div class="erp-section-dot"></div><h3>سجل طلبات جواب نموذج 41</h3></div>',unsafe_allow_html=True)
        reqs=load_requests()
        if reqs:
            req_df=pd.DataFrame(reqs)
            display_cols=['request_number','tax_number','supplier_name','periods','request_date','records_count','total_tax','created_at']
            display_names={'request_number':'رقم الطلب','tax_number':'التسجيل الضريبي','supplier_name':'اسم المورد','periods':'الفترات','request_date':'تاريخ الطلب','records_count':'عدد السجلات','total_tax':'اجمالي الضريبة','created_at':'تاريخ الإنشاء'}
            show_df=req_df[[c for c in display_cols if c in req_df.columns]].rename(columns=display_names)
            st.dataframe(show_df,use_container_width=True)
            del_labels=[f"طلب #{r.get('request_number','')} — {r.get('supplier_name','')} — {r.get('tax_number','')}" for r in reqs]
            del_sel=st.multiselect("اختر طلبات للحذف",options=list(range(len(reqs))),format_func=lambda i:del_labels[i],key="del_req")
            if del_sel and st.button("🗑️ حذف الطلبات المحددة",key="del_req_btn",type="primary"):
                new_reqs=[r for i,r in enumerate(reqs) if i not in del_sel]
                save_requests(new_reqs)
                st.success(f"تم حذف {len(del_sel)} طلب");st.rerun()
        else:
            st.info("لا توجد طلبات مسجلة بعد")

# ====================== VAT ======================
elif page == "💰 القيمة المضافة":
    if not user_has_permission(page): st.error("لا تملك صلاحية الوصول");st.stop()
    sub=st.radio("vat",["📤 رفع وربط البيانات","🔍 استعلام بالسجل الضريبي","📅 استعلام بالفترة"],horizontal=True,label_visibility="collapsed")
    if sub=="📤 رفع وربط البيانات":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>رفع شيت القيمة المضافة</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-info"><strong>الأعمدة:</strong> م • رقم التسجيل الضريبي • اسم الممول • 20% قيمة مضافة • ضريبة الجدول</div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        up=st.file_uploader("ارفع شيت القيمة المضافة",type=['xlsx','xls'],key="vat_up",label_visibility="collapsed")
        if up:
            if 'vat_df' not in st.session_state or st.session_state.get('vat_file')!=up.name:
                st.session_state['vat_df']=read_vat_excel(up);st.session_state['vat_file']=up.name
            df=st.session_state['vat_df']
            st.markdown(f'<div style="margin:.5rem 0;padding:.6rem 1rem;border-radius:10px;background:rgba(0,184,148,.08);border:1px solid rgba(0,184,148,.15);color:#55efc4;font-size:.82rem;">✓ تم الرفع — {len(df)} سجل</div>',unsafe_allow_html=True)
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>حذف سطور (اختياري)</h3></div>',unsafe_allow_html=True)
            row_labels=[f"سطر {i+1} — {df.iloc[i].to_dict().get('اسم الممول','')}" for i in range(len(df))]
            rm=st.multiselect("اختر السطور اللي عايز تشيلها",options=list(range(len(df))),format_func=lambda i:row_labels[i],key="vat_rm")
            if rm:
                st.markdown(f'<div style="padding:.5rem 1rem;border-radius:8px;background:rgba(255,107,107,.08);border:1px solid rgba(255,107,107,.15);color:#ff6b6b;font-size:.8rem;margin:.5rem 0;">⚠ سيتم حذف {len(rm)} سطر — المتبقي: {len(df)-len(rm)}</div>',unsafe_allow_html=True)
                if st.button("🗑️ تطبيق الحذف",key="vat_arm"):
                    st.session_state['vat_df']=df.drop(rm).reset_index(drop=True);st.rerun()
            df=st.session_state['vat_df']
            st.dataframe(df,use_container_width=True,height=280)
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>ربط البيانات</h3></div>',unsafe_allow_html=True)
            c1,c2=st.columns(2)
            with c1: vpd=st.date_input("تاريخ المدفوعة",key="vat_pd")
            with c2: vpn=st.text_input("رقم المدفوعة",key="vat_pn")
            c3,c4=st.columns(2)
            with c3: vm=st.selectbox("شهر النموذج",range(1,13),format_func=lambda x:f"{x}-{MONTHS[x]}",key="vat_m")
            with c4: vy=st.selectbox("سنة النموذج",range(2020,2031),key="vat_y")
            if st.button("💾 حفظ وربط البيانات",key="vat_save",type="primary"):
                data=load_data(VAT_FILE)
                data.append({"id":str(uuid.uuid4()),"upload_date":datetime.now().isoformat(),"payment_date":vpd.isoformat(),"payment_number":vpn,"model_month":vm,"model_year":vy,"file_name":up.name,"records":df.to_dict('records')})
                save_data(VAT_FILE,data);st.success("تم الحفظ!");del st.session_state['vat_df']
        st.markdown('</div>',unsafe_allow_html=True)
    elif sub=="🔍 استعلام بالسجل الضريبي":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>استعلام بالسجل الضريبي</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        c1,c2,c3=st.columns(3)
        with c1: qt=st.text_input("رقم التسجيل الضريبي",key="vat_qt")
        with c2: qm=st.selectbox("الشهر",range(1,13),index=6,format_func=lambda x:f"{x}-{MONTHS[x]}",key="vat_qm")
        with c3: qy=st.selectbox("السنة",range(2020,2031),key="vat_qy")
        if st.button("🔍 بحث",key="vat_bt",type="primary"):
            if not qt.strip(): st.warning("أدخل رقم التسجيل الضريبي")
            else:
                data=load_data(VAT_FILE);res=[]
                for rec in data:
                    if rec['model_month']==qm and rec['model_year']==qy:
                        for row in rec.get('records',[]):
                            if str(row.get('رقم التسجيل الضريبي','')).strip()==qt.strip():
                                rc=dict(row);rc['فترة الخصم']=f"{qm}/{qy}";rc['رقم المدفوعة']=rec.get('payment_number','');rc['تاريخ المدفوعة']=rec.get('payment_date','');res.append(rc)
                if res:
                    sn=res[0].get('اسم الممول','')
                    s1,s2,s3=st.columns(3)
                    with s1: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">اسم الممول</div><div class="erp-stat-value" style="font-size:.95rem">{sn}</div></div>',unsafe_allow_html=True)
                    with s2: st.markdown(f'<div class="erp-stat s-cyan"><div class="erp-stat-label">السجل الضريبي</div><div class="erp-stat-value" style="font-size:.95rem">{qt}</div></div>',unsafe_allow_html=True)
                    with s3: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الفترة</div><div class="erp-stat-value" style="font-size:1.1rem">{qm}/{qy}</div></div>',unsafe_allow_html=True)
                    tt=sum(_sf(r.get('ضريبة الجدول',0)) for r in res);tv=sum(_sf(r.get('20% قيمة مضافة',0)) for r in res)
                    s4,s5,s6=st.columns(3)
                    with s4: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">ضريبة الجدول</div><div class="erp-stat-value">{fmt(tt)}</div></div>',unsafe_allow_html=True)
                    with s5: st.markdown(f'<div class="erp-stat s-pink"><div class="erp-stat-label">20% قيمة مضافة</div><div class="erp-stat-value">{fmt(tv)}</div></div>',unsafe_allow_html=True)
                    with s6: st.markdown(f'<div class="erp-stat s-red"><div class="erp-stat-label">الإجمالي</div><div class="erp-stat-value">{fmt(tt+tv)}</div></div>',unsafe_allow_html=True)
                    st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>التفاصيل</h3></div>',unsafe_allow_html=True)
                    st.dataframe(pd.DataFrame(res),use_container_width=True)
                else: st.info("لم يتم العثور على نتائج")
        st.markdown('</div>',unsafe_allow_html=True)
    elif sub=="📅 استعلام بالفترة":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>استعلام بالفترة</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        qt=st.text_input("رقم التسجيل الضريبي",key="vat_pqt")
        periods_list=[(m,y) for y in range(2020,2031) for m in range(1,13)]
        period_labels={f"{m}/{y}":f"{m}-{MONTHS[m]} {y}" for m,y in periods_list}
        sel=st.multiselect("اختر الفترات",options=list(period_labels.keys()),format_func=lambda x:period_labels[x],default=[],key="vat_periods")
        if st.button("🔍 بحث",key="vat_bp",type="primary"):
            if not sel: st.warning("اختر فترة واحدة على الأقل")
            else:
                periods=set()
                for s in sel:
                    parts=s.split('/')
                    periods.add((int(parts[0]),int(parts[1])))
                data=load_data(VAT_FILE);res=[]
                for rec in data:
                    if (rec['model_month'],rec['model_year']) in periods:
                        for row in rec.get('records',[]):
                            if qt.strip():
                                if str(row.get('رقم التسجيل الضريبي','')).strip()==qt.strip():
                                    rc=dict(row);rc['الفترة']=f"{rec['model_month']}/{rec['model_year']}";rc['رقم المدفوعة']=rec.get('payment_number','');rc['تاريخ المدفوعة']=rec.get('payment_date','');res.append(rc)
                            else:
                                rc=dict(row);rc['الفترة']=f"{rec['model_month']}/{rec['model_year']}";rc['رقم المدفوعة']=rec.get('payment_number','');rc['تاريخ المدفوعة']=rec.get('payment_date','');res.append(rc)
                st.session_state['vat_period_res']=res
                st.session_state['vat_period_sel']=sel
                st.session_state['vat_period_qt']=qt
        res=st.session_state.get('vat_period_res',[])
        sel=st.session_state.get('vat_period_sel',[])
        qt_val=st.session_state.get('vat_period_qt','')
        if res:
            tt=sum(_sf(r.get('ضريبة الجدول',0)) for r in res);tv=sum(_sf(r.get('20% قيمة مضافة',0)) for r in res)
            st.success(f"{len(res)} سجل")
            s1,s2,s3=st.columns(3)
            with s1: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">ضريبة الجدول</div><div class="erp-stat-value">{fmt(tt)}</div></div>',unsafe_allow_html=True)
            with s2: st.markdown(f'<div class="erp-stat s-pink"><div class="erp-stat-label">20% قيمة مضافة</div><div class="erp-stat-value">{fmt(tv)}</div></div>',unsafe_allow_html=True)
            with s3: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الإجمالي</div><div class="erp-stat-value">{fmt(tt+tv)}</div></div>',unsafe_allow_html=True)
            st.dataframe(pd.DataFrame(res),use_container_width=True)
            st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>تصدير جواب القيمة المضافة</h3></div>',unsafe_allow_html=True)
            export_card='<div class="erp-card"><div class="erp-card-header"><div class="erp-card-icon" style="background:linear-gradient(135deg,rgba(253,203,110,.15),rgba(253,203,110,.03));">📝</div><div><h3>إعدادات التصدير</h3></div></div></div>'
            st.markdown(export_card,unsafe_allow_html=True)
            c1,c2=st.columns(2)
            with c1: vexport_date=st.date_input("تاريخ تقديم طلب الشهادة",key="vat_ed")
            with c2:
                vreqs=load_requests()
                vnext_num=1
                if vreqs:
                    vnums=[r.get('request_number',0) for r in vreqs]
                    vnext_num=max(vnums)+1 if vnums else 1
                vreq_num=st.number_input("رقم الطلب",min_value=1,value=vnext_num,step=1,key="vat_rn")
            vsn=res[0].get('اسم الممول','')
            vtn=str(res[0].get('رقم التسجيل الضريبي','')).strip()
            st.markdown(f'<div style="padding:.5rem 1rem;border-radius:10px;background:rgba(108,92,231,.06);border:1px solid rgba(108,92,231,.1);font-size:.82rem;margin:.5rem 0;">المورد: <strong>{vsn}</strong> — التسجيل الضريبي: <strong>{vtn}</strong></div>',unsafe_allow_html=True)
            if st.button("📄 إنشاء جواب القيمة المضافة",key="vat_export",type="primary"):
                vexport_date_formatted=vexport_date.strftime("%d/%m/%Y")
                vrequest_date_formatted=vexport_date.strftime("%d/%m/%Y")
                with st.spinner("جاري إنشاء الملف..."):
                    vout_path,verr=gen_vat_word(vsn,vtn,res,vexport_date_formatted,vrequest_date_formatted,int(vreq_num))
                if verr:
                    st.error(f"خطأ: {verr}")
                else:
                    vreq_record={
                        "id":str(uuid.uuid4()),
                        "request_number":int(vreq_num),
                        "tax_number":vtn,
                        "supplier_name":vsn,
                        "periods":sorted(list(sel)),
                        "request_date":vexport_date.isoformat(),
                        "export_date":vexport_date_formatted,
                        "records_count":len(res),
                        "total_tax":tt+tv,
                        "created_at":datetime.now().isoformat(),
                        "file_name":os.path.basename(vout_path)
                    }
                    vreqs=load_requests()
                    vreqs.append(vreq_record)
                    save_requests(vreqs)
                    st.success(f"تم إنشاء الملف — رقم الطلب: {vreq_num}")
                    with open(vout_path,'rb') as f:
                        st.download_button("📥 تحميل الجواب",data=f.read(),file_name=os.path.basename(vout_path),mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",type="primary")
        else: st.info("لا توجد سجلات في هذه الفترات")
        st.markdown('</div>',unsafe_allow_html=True)

# ====================== MARKET ======================
elif page=="🛒 فواتير الماركت":
    if not user_has_permission(page): st.error("لا تملك صلاحية الوصول");st.stop()
    st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>فواتير الماركت — إنشاء Template Portal</h3></div>',unsafe_allow_html=True)
    if 'mkt_step' not in st.session_state: st.session_state['mkt_step']=1

    steps_labels={1:"1️⃣ رفع Detailed Receipt",2:"2️⃣ رفع Barcodes",3:"3️⃣ إنشاء Template Portal"}
    cs=st.session_state['mkt_step']
    step_html='<div style="display:flex;align-items:center;gap:.5rem;margin-bottom:1.5rem;">'
    for si,sl in steps_labels.items():
        if si<cs: step_html+=f'<div style="display:flex;align-items:center;gap:.4rem;"><div style="width:30px;height:30px;border-radius:50%;background:linear-gradient(135deg,#00b894,#55efc4);display:flex;align-items:center;justify-content:center;color:#fff;font-size:.75rem;font-weight:700;">✓</div><span style="color:#55efc4;font-size:.78rem;font-weight:600;">الخطوة {si}</span></div><div style="width:30px;height:2px;background:linear-gradient(90deg,#00b894,rgba(0,184,148,.2));border-radius:2px;"></div>'
        elif si==cs: step_html+=f'<div style="display:flex;align-items:center;gap:.4rem;"><div style="width:30px;height:30px;border-radius:50%;background:linear-gradient(135deg,#6c5ce7,#a29bfe);display:flex;align-items:center;justify-content:center;color:#fff;font-size:.75rem;font-weight:700;box-shadow:0 0 15px rgba(108,92,231,.4);">{si}</div><span style="color:#fff;font-size:.78rem;font-weight:700;">الخطوة {si}</span></div>'
        else: step_html+=f'<div style="display:flex;align-items:center;gap:.4rem;"><div style="width:30px;height:30px;border-radius:50%;background:rgba(255,255,255,.04);border:1px solid rgba(255,255,255,.08);display:flex;align-items:center;justify-content:center;color:rgba(255,255,255,.25);font-size:.75rem;font-weight:700;">{si}</div><span style="color:rgba(255,255,255,.25);font-size:.78rem;">الخطوة {si}</span></div>'
        if si<3: step_html+='<div style="width:30px;height:2px;background:rgba(255,255,255,.06);border-radius:2px;"></div>'
    step_html+='</div>'
    st.markdown(step_html,unsafe_allow_html=True)

    if cs==1:
        st.markdown("""<div class="erp-card"><div class="erp-card-header">
            <div class="erp-card-icon" style="background:linear-gradient(135deg,rgba(108,92,231,.15),rgba(108,92,231,.03));">📄</div>
            <div><h3>الخطوة 1 — Detailed Receipt</h3><p>B-C(الضريبة) • F-E(الخصم) • G(الكمية) • H-I(السعر) • S-J(اسم الصنف) • T(رقم الصنف)</p></div>
        </div></div>""",unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        up=st.file_uploader("ارفع Detailed Receipt",type=['xlsx','xls'],key="det_up",label_visibility="collapsed")
        if up:
            try:
                ddf=read_detailed_receipt(up)
                if ddf.empty: st.error("الملف فارغ")
                else:
                    st.markdown(f'<div style="margin:.5rem 0;padding:.6rem 1rem;border-radius:10px;background:rgba(0,184,148,.08);border:1px solid rgba(0,184,148,.15);color:#55efc4;font-size:.82rem;">✓ {len(ddf)} صنف</div>',unsafe_allow_html=True)
                    st.session_state['detail_df']=ddf
                    st.dataframe(ddf,use_container_width=True,height=300)
                    if st.button("التالي ←",key="next1",type="primary"):
                        st.session_state['mkt_step']=2;st.rerun()
            except Exception as e: st.error(f"خطأ: {e}")
        st.markdown('</div>',unsafe_allow_html=True)
    elif cs==2:
        st.markdown("""<div class="erp-card"><div class="erp-card-header">
            <div class="erp-card-icon" style="background:linear-gradient(135deg,rgba(0,206,201,.15),rgba(0,206,201,.03));">📊</div>
            <div><h3>الخطوة 2 — Barcodes</h3><p>العمود B: رقم الصنف الداخلي | العمود G: الباركود</p></div>
        </div></div>""",unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        up=st.file_uploader("ارفع Barcodes",type=['xlsx','xls'],key="bc_up",label_visibility="collapsed")
        if up:
            try:
                bmap=read_barcodes(up)
                if not bmap: st.error("لا توجد بيانات")
                else:
                    st.markdown(f'<div style="margin:.5rem 0;padding:.6rem 1rem;border-radius:10px;background:rgba(0,184,148,.08);border:1px solid rgba(0,184,148,.15);color:#55efc4;font-size:.82rem;">✓ {len(bmap)} صنف مرتبط</div>',unsafe_allow_html=True)
                    st.session_state['barcode_map']=bmap
                    st.dataframe(pd.DataFrame(list(bmap.items())[:10],columns=['الكود الداخلي','الباركود']),use_container_width=True)
                    c1,c2=st.columns(2)
                    with c1:
                        if st.button("← السابق",key="prev2"):
                            st.session_state['mkt_step']=1;st.rerun()
                    with c2:
                        if st.button("التالي ←",key="next2",type="primary"):
                            st.session_state['mkt_step']=3;st.rerun()
            except Exception as e: st.error(f"خطأ: {e}")
        else:
            if st.button("← السابق",key="prev2b"):
                st.session_state['mkt_step']=1;st.rerun()
        st.markdown('</div>',unsafe_allow_html=True)
    elif cs==3:
        hd='detail_df' in st.session_state and not st.session_state['detail_df'].empty
        hb='barcode_map' in st.session_state and len(st.session_state.get('barcode_map',{}))>0
        if hd and hb:
            ddf=st.session_state['detail_df'];bmap=st.session_state['barcode_map']
            matched=sum(1 for _,r in ddf.iterrows() if str(r['internal_code']).strip() in bmap)
            unmatched=len(ddf)-matched;tc=sum(1 for _,r in ddf.iterrows() if float(r['tax'] or 0)>0)
            dt=sum(1 for _,r in ddf.iterrows() if float(r['discount'] or 0)>0 and float(r['tax'] or 0)>0)
            st.markdown("""<div class="erp-card"><div class="erp-card-header">
                <div class="erp-card-icon" style="background:linear-gradient(135deg,rgba(253,203,110,.15),rgba(253,203,110,.03));">⚙️</div>
                <div><h3>الخطوة 3 — ملخص البيانات</h3></div>
            </div></div>""",unsafe_allow_html=True)
            s1,s2,s3=st.columns(3)
            with s1: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">إجمالي الأصناف</div><div class="erp-stat-value">{len(ddf)}</div></div>',unsafe_allow_html=True)
            with s2: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">مرتبطة بالباركود</div><div class="erp-stat-value">{matched}</div></div>',unsafe_allow_html=True)
            with s3: st.markdown(f'<div class="erp-stat s-red"><div class="erp-stat-label">غير مرتبطه</div><div class="erp-stat-value">{unmatched}</div></div>',unsafe_allow_html=True)
            s4,s5=st.columns(2)
            with s4: st.markdown(f'<div class="erp-stat s-orange"><div class="erp-stat-label">بها ضريبة (÷1.14)</div><div class="erp-stat-value">{tc}</div></div>',unsafe_allow_html=True)
            with s5: st.markdown(f'<div class="erp-stat s-pink"><div class="erp-stat-label">خصم + ضريبة</div><div class="erp-stat-value">{dt}</div></div>',unsafe_allow_html=True)
            if unmatched>0:
                miss=[str(r['internal_code']).strip() for _,r in ddf.iterrows() if str(r['internal_code']).strip() not in bmap]
                st.warning(f"أكواد غير موجودة: {', '.join(miss[:15])}{'...' if len(miss)>15 else ''}")
            c1,c2=st.columns([1,2])
            with c1:
                if st.button("← السابق",key="prev3"):
                    st.session_state['mkt_step']=2;st.rerun()
            with c2:
                if st.button("🚀 إنشاء Template Portal",key="gen_p",type="primary"):
                    try:
                        out=gen_template(ddf,bmap);st.success("تم الإنشاء!")
                        st.dataframe(pd.read_excel(out,engine='openpyxl'),use_container_width=True,height=400)
                        out.seek(0)
                        st.download_button("📥 تحميل",data=out,file_name=f"Template_Portal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",type="primary")
                    except Exception as e: st.error(f"خطأ: {e}")
        else:
            miss=[]
            if not hd: miss.append("Detailed Receipt")
            if not hb: miss.append("Barcodes")
            st.markdown(f'<div class="erp-empty"><div class="erp-empty-icon">⏳</div><h3>بانتظار إكمال الخطوات</h3><p>يرجى رفع: {" + ".join(miss)}</p></div>',unsafe_allow_html=True)
            if st.button("← السابق",key="prev3b"):
                st.session_state['mkt_step']=2;st.rerun()

# ====================== PORTAL ELECTRONIC INVOICES ======================
elif page=="📄 Portal الفواتير الإلكترونية":
    if not user_has_permission(page): st.error("لا تملك صلاحية الوصول");st.stop()

    st.markdown(f"""<div class="erp-topbar"><div><h2>{page}</h2><p>إدارة فواتير الصادرة والواردة من بوابة الفواتير الإلكترونية</p></div>
<div class="erp-topbar-right"><a href="https://invoicing.eta.gov.eg/" target="_blank" style="background:linear-gradient(135deg,rgba(0,206,201,.18),rgba(108,92,231,.12));border:1px solid rgba(0,206,201,.35);border-radius:12px;padding:.5rem 1.2rem;color:#00cec9;font-size:.82rem;font-weight:700;text-decoration:none;cursor:pointer;transition:all .3s;display:inline-flex;align-items:center;gap:.5rem;">🔗 فتح بوابة الفواتير الإلكترونية</a></div></div>""", unsafe_allow_html=True)

    _tab1,_tab2,_tab3=st.tabs(["🔗 الربط","📤 الصادرة","📥 الوارد"])

    def _portal_dashboard(data,label,color_icon,label_type):
        if not data:
            st.info(f"لا توجد فواتير {label} بعد");return
        now=datetime.now()
        recent=[r for r in data if (now-datetime.fromisoformat(str(r.get('upload_date',''))[:19])).days<=30]
        total_recent=sum(r.get('records_count',len(r.get('records',[]))) for r in recent)
        valid_statuses=['مقبولة','مستلمة']
        invalid_statuses=['مرفوضة','ملغاة']
        n_valid=sum(1 for r in recent if r.get('status','') in valid_statuses)
        n_invalid=sum(1 for r in recent if r.get('status','') in invalid_statuses)
        n_total_batches=len(recent)
        total_all=sum(r.get('records_count',len(r.get('records',[]))) for r in data)
        s1,s2,s3,s4=st.columns(4)
        with s1: st.markdown(f'<div class="erp-stat {color_icon}"><div class="erp-stat-label">إجمالي {label} (30 يوم)</div><div class="erp-stat-value">{total_recent}</div><div class="erp-stat-sub">{n_total_batches} رفع</div></div>',unsafe_allow_html=True)
        with s2: st.markdown(f'<div class="erp-stat s-green"><div class="erp-stat-label">الفواتير الصحيحة</div><div class="erp-stat-value">{n_valid}</div><div class="erp-stat-sub">رفعات مقبولة</div></div>',unsafe_allow_html=True)
        with s3: st.markdown(f'<div class="erp-stat s-red"><div class="erp-stat-label">الفواتير الملغاة</div><div class="erp-stat-value">{n_invalid}</div><div class="erp-stat-sub">رفعات مرفوضة</div></div>',unsafe_allow_html=True)
        with s4: st.markdown(f'<div class="erp-stat s-blue"><div class="erp-stat-label">الإجمالي الكلي</div><div class="erp-stat-value">{total_all}</div><div class="erp-stat-sub">كل الفترات</div></div>',unsafe_allow_html=True)

    def _portal_filter_and_download(data,label_type):
        if not data: return
        import zipfile, tempfile
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>بحث وتصفية الفواتير</h3></div>',unsafe_allow_html=True)
        periods_list=[]
        for r in data:
            p=r.get('period','')
            if p and p not in periods_list: periods_list.append(p)

        c1,c2,c3,c4=st.columns(4)
        with c1:
            sel_period=st.selectbox("اختر الفترة (شهر/سنة)",options=["الكل"]+periods_list,key=f"fp_{label_type}")
        with c2:
            avail_dates=[]
            for r in data:
                ud=str(r.get('upload_date',''))[:10]
                if sel_period!="الكل" and r.get('period','')!=sel_period: continue
                if ud and ud not in avail_dates: avail_dates.append(ud)
            avail_dates=sorted(avail_dates,reverse=True)
            sel_date=st.selectbox("اختر يوم (تاريخ الرفع)",options=["الكل"]+avail_dates,key=f"fd_{label_type}")
        with c3:
            issue_dates=[]
            for r in data:
                if sel_period!="الكل" and r.get('period','')!=sel_period: continue
                for rec in r.get('records',[]):
                    idt=str(rec.get('تاريخ الإصدار',''))[:10]
                    if idt and idt not in issue_dates: issue_dates.append(idt)
            issue_dates=sorted(issue_dates,reverse=True)
            sel_issue_date=st.selectbox("اختر يوم (تاريخ الإصدار)",options=["الكل"]+issue_dates,key=f"fiss_{label_type}")
        with c4:
            all_statuses=set()
            for r in data:
                s=r.get('status','')
                if s: all_statuses.add(s)
            sel_status=st.selectbox("الحالة",options=["الكل","مقبولة","مستلمة","مرفوضة","مرسلة","ملغاة"]+sorted(all_statuses),key=f"fs_{label_type}")

        c4,c5=st.columns(2)
        with c4:
            all_types=set()
            for r in data:
                for rec in r.get('records',[]):
                    t=rec.get('نوع الفاتورة','')
                    if t: all_types.add(t)
            inv_type_list=sorted(all_types) if all_types else ["فاتورة بيع","إشعار دائن","إشعار مدين"]
            sel_inv_type=st.selectbox("نوع الفاتورة",options=["الكل"]+inv_type_list,key=f"ft_{label_type}")
        with c5:
            st.write("")

        filtered=[]
        for r in data:
            if sel_period!="الكل" and r.get('period','')!=sel_period: continue
            if sel_date!="الكل" and str(r.get('upload_date',''))[:10]!=sel_date: continue
            if sel_status!="الكل" and r.get('status','')!=sel_status: continue
            if sel_issue_date!="الكل":
                recs=r.get('records',[])
                matched_recs=[rec for rec in recs if str(rec.get('تاريخ الإصدار',''))[:10]==sel_issue_date]
                if not matched_recs: continue
            filtered.append(r)

        if not filtered:
            st.info("لا توجد فواتير تطابق البحث");return

        total_recs=sum(r.get('records_count',len(r.get('records',[]))) for r in filtered)
        st.success(f"تم العثور على {len(filtered)} رفع يحتوي على {total_recs} فاتورة")

        all_records=[]
        for rec in filtered:
            recs=rec.get('records',[])
            for r in recs:
                r['__period__']=rec.get('period','')
                r['__type__']=rec.get('invoice_type','')
                r['__status__']=rec.get('status','')
                r['__upload_date__']=str(rec.get('upload_date',''))[:10]
                r['__file_name__']=rec.get('file_name','')
            if sel_inv_type!="الكل":
                recs=[r for r in recs if r.get('نوع الفاتورة','')==sel_inv_type]
            if sel_issue_date!="الكل":
                recs=[r for r in recs if str(r.get('تاريخ الإصدار',''))[:10]==sel_issue_date]
            all_records.extend(recs)

        if st.button("📋 تفاصيل فواتير الفترة",key=f"detail_btn_{label_type}",type="primary"):
            st.session_state[f"show_details_{label_type}"]=True
            st.session_state[f"detail_records_{label_type}"]=all_records
            st.rerun()

        if st.session_state.get(f"show_details_{label_type}"):
            detail_recs=st.session_state.get(f"detail_records_{label_type}",[])
            st.markdown(f"""<div style="background:linear-gradient(135deg,rgba(108,92,231,.12),rgba(0,206,201,.08));border:2px solid rgba(108,92,231,.3);border-radius:16px;padding:1.5rem;margin:1rem 0;box-shadow:0 8px 32px rgba(108,92,231,.2);">
            <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:1rem;border-bottom:1px solid rgba(255,255,255,.06);padding-bottom:.8rem;">
                <h3 style="color:#fff;margin:0;font-size:1.1rem;">تفاصيل فواتير الفترة ({len(detail_recs)} فاتورة)</h3>
            </div>
            </div>""",unsafe_allow_html=True)
            for idx,r in enumerate(detail_recs):
                st_status=r.get('الحالة',r.get('__status__',''))
                if st_status in ['مقبولة','مستلمة']:
                    bulb_color='#55efc4'
                elif st_status in ['ملغاة','مرفوضة']:
                    bulb_color='#ff6b6b'
                else:
                    bulb_color='#fdcb6e'
                supplier=r.get('الطرف الآخر',r.get('اسم المصدر',r.get('اسم المستلم','-')))
                inv_type=r.get('نوع الفاتورة',r.get('__type__',''))
                inv_total=r.get('الإجمالي (بعد الضريبة)',r.get('الإجمالي',0))
                inv_date=r.get('تاريخ الإصدار','')
                uuid_val=r.get('UUID','—')
                st.markdown(f"""<div style="background:rgba(30,30,56,.8);border:1px solid rgba(255,255,255,.06);border-radius:10px;padding:.7rem 1rem;margin-bottom:.5rem;">
                <div style="display:flex;justify-content:space-between;align-items:center;">
                    <div>
                        <div style="color:#fff;font-weight:700;font-size:.88rem;">{supplier}</div>
                        <div style="display:flex;gap:.8rem;margin-top:.3rem;flex-wrap:wrap;">
                            <span style="color:var(--text2);font-size:.72rem;">📋 {inv_type}</span>
                            <span style="color:var(--text2);font-size:.72rem;">📅 {inv_date}</span>
                            <span style="color:#a29bfe;font-size:.78rem;font-weight:600;">الإجمالي: {fmt(inv_total)}</span>
                        </div>
                        <div style="color:var(--text2);font-size:.68rem;margin-top:.2rem;font-family:monospace;">UUID: {uuid_val}</div>
                    </div>
                    <div style="display:flex;align-items:center;gap:.4rem;">
                        <div style="width:10px;height:10px;border-radius:50%;background:{bulb_color};box-shadow:0 0 8px {bulb_color};"></div>
                        <span style="color:{bulb_color};font-size:.75rem;font-weight:600;">{st_status}</span>
                    </div>
                </div>
            </div>""",unsafe_allow_html=True)
            if st.button("✕ إغلاق",key=f"close_details_{label_type}"):
                st.session_state[f"show_details_{label_type}"]=False
                st.rerun()

        if all_records:
            all_records=_fix_vat_in_records(all_records)
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>تحميل جماعي</h3></div>',unsafe_allow_html=True)
            gc1,gc2,gc3=st.columns(3)
            with gc1:
                all_df=pd.DataFrame(all_records)
                excel_buf=BytesIO()
                all_df.to_excel(excel_buf,index=False,engine='xlsxwriter')
                excel_buf.seek(0)
                st.download_button(f"📊 تحميل Excel ({len(all_records)} فاتورة)",data=excel_buf.getvalue(),file_name=f"all_{label_type}_{sel_period.replace('/','_') if sel_period!='الكل' else 'all'}.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",key=f"dl_all_{label_type}",type="primary")
            with gc2:
                pdf_buf=_generate_pdf_for_records(all_records,f"فواتير {'الصادرة' if label_type=='out' else 'الواردة'} — {sel_period if sel_period!='الكل' else 'الكل'}")
                st.download_button(f"📄 تحميل PDF ({len(all_records)} فاتورة)",data=pdf_buf.getvalue(),file_name=f"invoices_{label_type}_{sel_period.replace('/','_') if sel_period!='الكل' else 'all'}.pdf",mime="application/pdf",key=f"dl_all_pdf_{label_type}")
            with gc3:
                zip_buf=BytesIO()
                with zipfile.ZipFile(zip_buf,'w',zipfile.ZIP_DEFLATED) as zf:
                    for ri,r in enumerate(all_records):
                        supplier=r.get('الطرف الآخر',r.get('اسم المصدر',r.get('اسم المستلم','')))
                        safe_supplier=''.join(c if c.isalnum() or c in '_-' else '_' for c in str(supplier))[:30]
                        pdf_inv=_generate_pdf_for_records([r],f"فاتورة #{ri+1} — {supplier}")
                        zf.writestr(f"فاتورة_{ri+1}_{safe_supplier}.pdf",pdf_inv.getvalue())
                    all_df2=pd.DataFrame(all_records)
                    excel_buf2=BytesIO()
                    all_df2.to_excel(excel_buf2,index=False,engine='xlsxwriter')
                    excel_buf2.seek(0)
                    zf.writestr("ملخص.xlsx",excel_buf2.getvalue())
                zip_buf.seek(0)
                st.download_button(f"📦 تحميل مضغوط ({len(all_records)} فاتورة PDF+Excel)",data=zip_buf.getvalue(),file_name=f"invoices_{label_type}_{sel_period.replace('/','_') if sel_period!='الكل' else 'all'}.zip",mime="application/zip",key=f"dl_zip_{label_type}")

    with _tab1:
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>ربط مباشر مع بوابة الفواتير الإلكترونية</h3></div>',unsafe_allow_html=True)

        eta_cid=os.environ.get("ETA_CLIENT_ID","")
        eta_csec=os.environ.get("ETA_CLIENT_SECRET","")
        if not eta_cid:
            try:
                eta_cid=st.secrets.get("ETA_CLIENT_ID","")
                eta_csec=st.secrets.get("ETA_CLIENT_SECRET","")
            except:
                pass

        if eta_cid and not st.session_state.get("eta_token"):
            with st.spinner("جاري الاتصال التلقائي بالبورتال..."):
                token,err=eta_login(eta_cid,eta_csec)
            if err:
                st.error(f"❌ فشل الاتصال التلقائي: {err}")
                st.info("يمكنك إدخال البيانات يدوياً من النموذج أدناه")
            else:
                st.session_state["eta_token"]=token
                st.session_state["eta_client_id"]=eta_cid
                st.session_state["eta_client_secret"]=eta_csec
                st.success("✅ تم الاتصال التلقائي بالبورتال!")
                st.rerun()

        if not eta_cid:
            st.markdown("""<div style="padding:.8rem 1rem;border-radius:10px;background:rgba(253,203,110,.06);border:1px solid rgba(253,203,110,.15);margin-bottom:1rem;font-size:.78rem;color:#fdcb6e;">
                💡 أدخل بيانات الاتصال يدوياً من البورتال
            </div>""",unsafe_allow_html=True)
            with st.form("eta_credentials_form"):
                c1,c2=st.columns(2)
                with c1:
                    eta_client_id_input=st.text_input("Client ID",value=st.session_state.get("eta_client_id",""),placeholder="أدخل Client ID")
                with c2:
                    eta_client_secret_input=st.text_input("Client Secret",value=st.session_state.get("eta_client_secret",""),type="password",placeholder="أدخل Client Secret")
                submitted=st.form_submit_button("🔐 الاتصال بالبورتال",type="primary")
            if submitted:
                if not eta_client_id_input or not eta_client_secret_input:
                    st.error("أدخل Client ID و Client Secret")
                else:
                    with st.spinner("جاري الاتصال بالبورتال..."):
                        token,err=eta_login(eta_client_id_input,eta_client_secret_input)
                    if err:
                        st.error(f"❌ فشل الاتصال: {err}")
                    else:
                        st.session_state["eta_token"]=token
                        st.session_state["eta_client_id"]=eta_client_id_input
                        st.session_state["eta_client_secret"]=eta_client_secret_input
                        st.success("✅ تم الاتصال بالبورتال بنجاح!")
                        st.rerun()

        if st.session_state.get("eta_token"):
            st.markdown('<div style="padding:.6rem 1rem;border-radius:10px;background:rgba(85,239,196,.06);border:1px solid rgba(85,239,196,.15);color:#55efc4;font-size:.82rem;margin:.5rem 0;">✅ متصلاً بالبورتال — جاهز لجلب الفواتير</div>',unsafe_allow_html=True)

            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>جلب الفواتير من البورتال</h3></div>',unsafe_allow_html=True)

            st.markdown("""<div style="padding:.6rem 1rem;border-radius:10px;background:rgba(116,185,255,.06);border:1px solid rgba(116,185,255,.15);color:#74b9ff;font-size:.78rem;margin-bottom:1rem;">
                📌 الحد الأقصى للفترة هو 30 يوم — هجلب الفواتير ونحفظها تلقائياً
            </div>""",unsafe_allow_html=True)

            today=datetime.now()
            default_from=today-timedelta(days=30)
            c1,c2,c3=st.columns(3)
            with c1:
                date_from=st.date_input("من تاريخ",value=default_from,key="eta_from")
            with c2:
                date_to=st.date_input("إلى تاريخ",value=today,key="eta_to")
            with c3:
                st.write("");st.write("")
                fetch_sent=st.checkbox("📥 جلب الصادرة",value=True,key="eta_fetch_sent")
                fetch_received=st.checkbox("📤 جلب الوارد",value=True,key="eta_fetch_received")

            if st.button("🔄 جلب الفواتير من البورتال",type="primary",key="eta_fetch_btn"):
                token=st.session_state["eta_token"]
                total_fetched=0
                errors=[]

                if fetch_sent:
                    with st.spinner("جاري جلب الفواتير الصادرة..."):
                        docs,err=eta_search_docs(token,"Sent",date_from,date_to)
                    if err:
                        errors.append(f"الصادرة: {err}")
                    else:
                        records=[eta_doc_to_record(d,"Sent") for d in docs]
                        out_data=load_data(PORTAL_OUT_FILE)
                        for _,(rec,meta) in enumerate(records):
                            meta["records"]=[rec]
                            meta["records_count"]=1
                            out_data.append(meta)
                        save_data(PORTAL_OUT_FILE,out_data)
                        total_fetched+=len(records)
                        st.success(f"✅ تم جلب {len(records)} فاتورة صادرة")

                if fetch_received:
                    with st.spinner("جاري جلب فواتير الوارد..."):
                        docs,err=eta_search_docs(token,"Received",date_from,date_to)
                    if err:
                        errors.append(f"الوارد: {err}")
                    else:
                        records=[eta_doc_to_record(d,"Received") for d in docs]
                        in_data=load_data(PORTAL_IN_FILE)
                        for _,(rec,meta) in enumerate(records):
                            meta["records"]=[rec]
                            meta["records_count"]=1
                            in_data.append(meta)
                        save_data(PORTAL_IN_FILE,in_data)
                        total_fetched+=len(records)
                        st.success(f"✅ تم جلب {len(records)} فاتورة وارد")

                if errors:
                    for e in errors:
                        st.error(f"⚠️ {e}")
                if total_fetched>0:
                    st.success(f"تم جلب {total_fetched} فاتورة بنجاح!")

            if st.button("🚪 قطع الاتصال",key="eta_disconnect"):
                for k in ["eta_token","eta_client_id","eta_client_secret"]:
                    if k in st.session_state: del st.session_state[k]
                st.success("تم قطع الاتصال");st.rerun()
        else:
            st.markdown("""<div class="erp-card" style="text-align:center;padding:2rem;">
                <div style="font-size:3rem;margin-bottom:1rem;">🔗</div>
                <h3 style="color:#fff;margin:0;">لم تتم بعد الاتصال بالبورتال</h3>
                <p style="color:var(--text2);font-size:.85rem;margin:.5rem 0;">أدخل بياناتك من البورتال المصري للاتصال وجلب الفواتير تلقائياً</p>
            </div>""",unsafe_allow_html=True)

    with _tab2:
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>فواتير الصادرة</h3></div>',unsafe_allow_html=True)

        out_data=load_data(PORTAL_OUT_FILE)
        _portal_dashboard(out_data,"الصادرة","s-orange","out")

        _portal_filter_and_download(out_data,"out")

        if out_data:
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>تحميل حزمة فواتير الصادرة</h3></div>',unsafe_allow_html=True)
            st.markdown('<div class="erp-card">',unsafe_allow_html=True)
            periods_list=[]
            for r in out_data:
                p=r.get('period','')
                if p and p not in periods_list: periods_list.append(p)
            periods_list=sorted(periods_list,reverse=True)
            bc1,bc2,bc3,bc4=st.columns(4)
            with bc1:
                dl_period=st.selectbox("اختر الفترة",options=["الكل"]+periods_list,key="out_dl_period")
            with bc2:
                avail_dates_out=[]
                for r in out_data:
                    ud=str(r.get('upload_date',''))[:10]
                    if dl_period!="الكل" and r.get('period','')!=dl_period: continue
                    if ud and ud not in avail_dates_out: avail_dates_out.append(ud)
                avail_dates_out=sorted(avail_dates_out,reverse=True)
                dl_day=st.selectbox("اختر يوم",options=["الكل"]+avail_dates_out,key="out_dl_day")
            with bc3:
                dl_type=st.selectbox("نوع الفاتورة",["الكل","فاتورة بيع","إشعار دائن","إشعار مدين"],key="out_dl_type")
            with bc4:
                st.write("");st.write("")
            filtered_dl=[]
            for r in out_data:
                if dl_period!="الكل" and r.get('period','')!=dl_period: continue
                if dl_day!="الكل" and str(r.get('upload_date',''))[:10]!=dl_day: continue
                if dl_type!="الكل" and r.get('invoice_type','')!=dl_type: continue
                for rec in r.get('records',[]):
                    filtered_dl.append(rec)
            if filtered_dl:
                st.success(f"{len(filtered_dl)} فاتورة جاهزة للتحميل")
                if dl_day!="الكل":
                    dl_label=f"يوم {dl_day}"
                elif dl_period!="الكل":
                    dl_label=f"فترة {dl_period}"
                else:
                    dl_label="الكل"
                if st.button("📥 تحميل الحزمة",key="out_dl_btn",type="primary"):
                    import zipfile,time
                    token=st.session_state.get("eta_token","")
                    zip_buf=BytesIO()
                    progress=st.progress(0,text="جاري تحميل الفواتير...")
                    with zipfile.ZipFile(zip_buf,'w',zipfile.ZIP_DEFLATED) as zf:
                        for ri,rec in enumerate(filtered_dl):
                            progress.progress((ri+1)/len(filtered_dl),text=f"فاتورة {ri+1}/{len(filtered_dl)}")
                            uuid_val=rec.get('UUID','')
                            supplier=rec.get('الطرف الآخر',rec.get('اسم المصدر',rec.get('اسم المستلم',f'inv_{ri+1}')))
                            safe_supplier=''.join(c if c.isalnum() or c in '_-' else '_' for c in str(supplier))[:30]
                            if token and uuid_val:
                                pdf_buf,err=eta_get_document_pdf(token,uuid_val)
                                if not err and pdf_buf:
                                    zf.writestr(f"فاتورة_{ri+1}_{safe_supplier}.pdf",pdf_buf.getvalue())
                                    time.sleep(2.1)
                                    continue
                            pdf_inv=_generate_pdf_for_records([rec],f"فاتورة #{ri+1} — {supplier}")
                            zf.writestr(f"فاتورة_{ri+1}_{safe_supplier}.pdf",pdf_inv.getvalue())
                        df_bundle=pd.DataFrame(_fix_vat_in_records(filtered_dl))
                        buf=BytesIO()
                        df_bundle.to_excel(buf,index=False,engine='xlsxwriter')
                        buf.seek(0)
                        zf.writestr("ملخص.xlsx",buf.getvalue())
                    progress.empty()
                    zip_buf.seek(0)
                    st.download_button("📦 تحميل الملف المضغوط",data=zip_buf.getvalue(),file_name=f"fawatir_sadira_{dl_label.replace('/','_')}.zip",mime="application/zip",key="out_zip_dl")
            else:
                st.info("لا توجد فواتير تطابق الاختيار")
            st.markdown('</div>',unsafe_allow_html=True)

            if st.button("🗑️ حذف جميع الفواتير الصادرة",key="del_all_out",type="secondary"):
                save_data(PORTAL_OUT_FILE,[]);st.success("تم الحذف");st.rerun()

    with _tab3:
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>فواتير الوارد</h3></div>',unsafe_allow_html=True)

        in_data=load_data(PORTAL_IN_FILE)
        _portal_dashboard(in_data,"الواردة","s-cyan","in")

        _portal_filter_and_download(in_data,"in")

        if in_data:
            st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>تحميل حزمة فواتير الوارد</h3></div>',unsafe_allow_html=True)
            st.markdown('<div class="erp-card">',unsafe_allow_html=True)
            periods_list=[]
            for r in in_data:
                p=r.get('period','')
                if p and p not in periods_list: periods_list.append(p)
            periods_list=sorted(periods_list,reverse=True)
            bc1,bc2,bc3,bc4=st.columns(4)
            with bc1:
                dl_period=st.selectbox("اختر الفترة",options=["الكل"]+periods_list,key="in_dl_period")
            with bc2:
                avail_dates_in=[]
                for r in in_data:
                    ud=str(r.get('upload_date',''))[:10]
                    if dl_period!="الكل" and r.get('period','')!=dl_period: continue
                    if ud and ud not in avail_dates_in: avail_dates_in.append(ud)
                avail_dates_in=sorted(avail_dates_in,reverse=True)
                dl_day=st.selectbox("اختر يوم",options=["الكل"]+avail_dates_in,key="in_dl_day")
            with bc3:
                dl_type=st.selectbox("نوع الفاتورة",["الكل","فاتورة شراء","إشعار دائن وارد","إشعار مدين وارد"],key="in_dl_type")
            with bc4:
                st.write("");st.write("")
            filtered_dl=[]
            for r in in_data:
                if dl_period!="الكل" and r.get('period','')!=dl_period: continue
                if dl_day!="الكل" and str(r.get('upload_date',''))[:10]!=dl_day: continue
                if dl_type!="الكل" and r.get('invoice_type','')!=dl_type: continue
                for rec in r.get('records',[]):
                    filtered_dl.append(rec)
            if filtered_dl:
                st.success(f"{len(filtered_dl)} فاتورة جاهزة للتحميل")
                if dl_day!="الكل":
                    dl_label=f"يوم {dl_day}"
                elif dl_period!="الكل":
                    dl_label=f"فترة {dl_period}"
                else:
                    dl_label="الكل"
                if st.button("📥 تحميل الحزمة",key="in_dl_btn",type="primary"):
                    import zipfile,time
                    token=st.session_state.get("eta_token","")
                    zip_buf=BytesIO()
                    progress=st.progress(0,text="جاري تحميل الفواتير...")
                    with zipfile.ZipFile(zip_buf,'w',zipfile.ZIP_DEFLATED) as zf:
                        for ri,rec in enumerate(filtered_dl):
                            progress.progress((ri+1)/len(filtered_dl),text=f"فاتورة {ri+1}/{len(filtered_dl)}")
                            uuid_val=rec.get('UUID','')
                            supplier=rec.get('الطرف الآخر',rec.get('اسم المصدر',rec.get('اسم المستلم',f'inv_{ri+1}')))
                            safe_supplier=''.join(c if c.isalnum() or c in '_-' else '_' for c in str(supplier))[:30]
                            if token and uuid_val:
                                pdf_buf,err=eta_get_document_pdf(token,uuid_val)
                                if not err and pdf_buf:
                                    zf.writestr(f"فاتورة_{ri+1}_{safe_supplier}.pdf",pdf_buf.getvalue())
                                    time.sleep(2.1)
                                    continue
                            pdf_inv=_generate_pdf_for_records([rec],f"فاتورة #{ri+1} — {supplier}")
                            zf.writestr(f"فاتورة_{ri+1}_{safe_supplier}.pdf",pdf_inv.getvalue())
                        df_bundle=pd.DataFrame(_fix_vat_in_records(filtered_dl))
                        buf=BytesIO()
                        df_bundle.to_excel(buf,index=False,engine='xlsxwriter')
                        buf.seek(0)
                        zf.writestr("ملخص.xlsx",buf.getvalue())
                    progress.empty()
                    zip_buf.seek(0)
                    st.download_button("📦 تحميل الملف المضغوط",data=zip_buf.getvalue(),file_name=f"fawatir_warida_{dl_label.replace('/','_')}.zip",mime="application/zip",key="in_zip_dl")
            else:
                st.info("لا توجد فواتير تطابق الاختيار")
            st.markdown('</div>',unsafe_allow_html=True)

            if st.button("🗑️ حذف جميع فواتير الوارد",key="del_all_in",type="secondary"):
                save_data(PORTAL_IN_FILE,[]);st.success("تم الحذف");st.rerun()

# ====================== الاستعلام عن الأكواد ======================
elif page=="🏷️ الاستعلام عن الأكواد":
    if not user_has_permission(page): st.error("لا تملك صلاحية الوصول");st.stop()

    st.markdown(f"""<div class="erp-topbar"><div><h2>{page}</h2><p>بحث واستعلام عن أكواد الأصناف في الفواتير الإلكترونية</p></div>
<div class="erp-topbar-right"><span class="erp-badge">🏷️ أكواد</span></div></div>""", unsafe_allow_html=True)

    codes_db=load_codes_db()
    st.markdown(f"""<div style="padding:.6rem 1rem;border-radius:10px;background:rgba(116,185,255,.06);border:1px solid rgba(116,185,255,.15);color:#74b9ff;font-size:.82rem;margin-bottom:1rem;">
        📊 قاعدة الأكواد تحتوي على <strong>{len(codes_db)}</strong> صنف
    </div>""",unsafe_allow_html=True)

    out_data=load_data(PORTAL_OUT_FILE)
    in_data=load_data(PORTAL_IN_FILE)
    all_uuids=[]
    for rec in out_data:
        for r in rec.get('records',[]):
            uid=r.get('UUID','')
            if uid: all_uuids.append({'uuid':uid,'direction':'out','name':r.get('الطرف الآخر','')})
    for rec in in_data:
        for r in rec.get('records',[]):
            uid=r.get('UUID','')
            if uid: all_uuids.append({'uuid':uid,'direction':'in','name':r.get('الطرف الآخر','')})
    fetched_uuids=set(c.get('uuid','') for c in codes_db)
    unfetched=[u for u in all_uuids if u['uuid'] not in fetched_uuids]

    c1,c2=st.columns(2)
    with c1:
        st.markdown(f"""<div style="padding:.6rem 1rem;border-radius:10px;background:rgba(85,239,196,.06);border:1px solid rgba(85,239,196,.15);color:#55efc4;font-size:.82rem;">
            📦 إجمالي الفواتير: <strong>{len(all_uuids)}</strong> | تم استخراج الأكواد: <strong>{len(fetched_uuids)}</strong> | متبقي: <strong>{len(unfetched)}</strong>
        </div>""",unsafe_allow_html=True)
    with c2:
        if st.button("🔄 تحديث الأكواد من البورتال",key="refresh_codes",type="primary"):
            token=st.session_state.get("eta_token","")
            if not token:
                st.error("يجب الاتصال بالبورتال أولاً من تاب Portal الفواتير الإلكترونية")
            elif not unfetched:
                st.success("جميع الفواتير تم استخراج أكوادها بالفعل!")
            else:
                progress=st.progress(0,text=f"جاري استخراج الأكواد من {len(unfetched)} فاتورة...")
                new_codes=[]
                errors=0
                for idx,u in enumerate(unfetched):
                    progress.progress((idx+1)/len(unfetched),text=f"فاتورة {idx+1}/{len(unfetched)} — {u['uuid'][:12]}...")
                    doc,err=eta_get_document_details(token,u['uuid'])
                    if err or not doc:
                        errors+=1
                        continue
                    document=doc.get('document',{})
                    invoice_lines=document.get('invoiceLines',[])
                    for line in invoice_lines:
                        item_code=line.get('itemCode','')
                        item_desc=line.get('description','')
                        internal_code=line.get('internalCode','')
                        item_type=line.get('itemType','')
                        quantity=line.get('quantity',0)
                        unit_type=line.get('unitType','')
                        unit_val=line.get('unitValue',{})
                        unit_price=unit_val.get('amountEGP',0) if isinstance(unit_val,dict) else 0
                        sales_total=line.get('salesTotal',0)
                        new_codes.append({
                            'uuid':u['uuid'],'direction':'صادر' if u['direction']=='out' else 'وارد',
                            'counterparty':u['name'],'itemCode':item_code,'internalCode':internal_code,
                            'description':item_desc,'itemType':item_type,
                            'quantity':quantity,'unitType':unit_type,'unitPrice':unit_price,'salesTotal':sales_total
                        })
                    time.sleep(2.1)
                if new_codes:
                    codes_db.extend(new_codes)
                    codes_db_unique=[]
                    seen=set()
                    for c in codes_db:
                        key=(c.get('uuid',''),c.get('itemCode',''),c.get('description',''))
                        if key not in seen:
                            seen.add(key)
                            codes_db_unique.append(c)
                    save_codes_db(codes_db_unique)
                    st.success(f"تم استخراج {len(new_codes)} صنف من {len(unfetched)-errors} فاتورة (أخطاء: {errors})")
                    st.rerun()
                else:
                    st.warning("لم يتم العثور على أصناف في الفواتير المدروسة")
                progress.empty()

    st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>بحث عن صنف</h3></div>',unsafe_allow_html=True)
    st.markdown('<div class="erp-card">',unsafe_allow_html=True)
    search_query=st.text_input("اكتب كلمة مفتاحية (اسم الصنف أو الوصف أو الكود)",key="codes_search",placeholder="مثال: فلاش، منظف، أرز...")
    if st.button("🔍 بحث",key="codes_search_btn",type="primary"):
        if not search_query.strip():
            st.warning("أدخل كلمة للبحث")
        elif not codes_db:
            st.info("قاعدة الأكواد فاضية — اضغط تحديث الأكواد أولاً")
        else:
            q=search_query.strip()
            results=[]
            for c in codes_db:
                score=0
                desc=str(c.get('description','')).lower()
                code=str(c.get('itemCode','')).lower()
                ic=str(c.get('internalCode','')).lower()
                name=str(c.get('counterparty','')).lower()
                if q.lower() in desc: score=max(score,10)
                if q.lower() in code: score=max(score,8)
                if q.lower() in ic: score=max(score,8)
                if q.lower() in name: score=max(score,5)
                q_words=q.lower().split()
                for w in q_words:
                    if w in desc: score+=3
                    if w in code: score+=2
                    if w in ic: score+=2
                if score>0:
                    results.append((score,c))
            results.sort(key=lambda x:x[0],reverse=True)
            if results:
                st.success(f"تم العثور على {len(results)} نتيجة")
                seen_codes=set()
                for score,c in results:
                    item_key=(c.get('itemCode',''),c.get('description',''))
                    if item_key in seen_codes: continue
                    seen_codes.add(item_key)
                    dir_color='#00cec9' if c.get('direction')=='وارد' else '#a29bfe'
                    dir_label=c.get('direction','')
                    st.markdown(f"""<div class="erp-card" style="margin-bottom:.6rem;border-left:3px solid {dir_color};">
                        <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                            <div>
                                <div style="color:#fff;font-weight:700;font-size:.92rem;">{c.get('description','—')}</div>
                                <div style="display:flex;gap:.8rem;margin-top:.3rem;flex-wrap:wrap;">
                                    <span style="color:#00cec9;font-size:.78rem;font-weight:600;">📋 كود: {c.get('itemCode','—')}</span>
                                    <span style="color:#fdcb6e;font-size:.75rem;">رقم داخلي: {c.get('internalCode','—')}</span>
                                    <span style="color:{dir_color};font-size:.72rem;font-weight:600;">{dir_label}</span>
                                </div>
                                <div style="color:var(--text2);font-size:.72rem;margin-top:.2rem;">المورد/العميل: {c.get('counterparty','—')}</div>
                            </div>
                            <div style="text-align:left;">
                                <div style="color:#55efc4;font-weight:700;font-size:.85rem;">{fmt(c.get('salesTotal',0))}</div>
                                <div style="color:var(--text2);font-size:.68rem;">الكمية: {c.get('quantity',0)} {c.get('unitType','')}</div>
                            </div>
                        </div>
                    </div>""",unsafe_allow_html=True)
            else:
                st.info(f"لم يتم العثور على نتائج لكلمة: {q}")
    st.markdown('</div>',unsafe_allow_html=True)

    if codes_db:
        st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>جميع الأكواد</h3></div>',unsafe_allow_html=True)
        codes_df=pd.DataFrame([{'الكود':c.get('itemCode',''),'الكود الداخلي':c.get('internalCode',''),'الوصف':c.get('description',''),
            'الاتجاه':c.get('direction',''),'المورد/العميل':c.get('counterparty',''),'الكمية':c.get('quantity',0),
            'وحدة القياس':c.get('unitType',''),'سعر الوحدة':c.get('unitPrice',0),'الإجمالي':c.get('salesTotal',0)} for c in codes_db])
        st.dataframe(codes_df,use_container_width=True,height=400)
        excel_buf=BytesIO()
        codes_df.to_excel(excel_buf,index=False,engine='xlsxwriter')
        excel_buf.seek(0)
        st.download_button("📊 تحميل جميع الأكواد",data=excel_buf.getvalue(),file_name="item_codes.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",key="dl_codes")

# ====================== الاستعلام عن ممول ======================
elif page=="🔍 الاستعلام عن ممول":
    if not user_has_permission(page): st.error("لا تملك صلاحية الوصول");st.stop()

    st.markdown(f"""<div class="erp-topbar"><div><h2>{page}</h2><p>البحث عن ممولين بالسجل الضريبي أو الاسم</p></div>
<div class="erp-topbar-right"><span class="erp-badge">🔍 بحث</span></div></div>""",unsafe_allow_html=True)

    def _build_suppliers_db():
        suppliers={}
        f41=load_data(FORM41_FILE)
        for rec in f41:
            for r in rec.get('records',[]):
                tax_num=str(r.get('رقم التسجيل الضريبي','')).strip()
                name=str(r.get('اسم الممول','')).strip()
                if tax_num and name:
                    if tax_num not in suppliers:
                        suppliers[tax_num]={'tax_number':tax_num,'names':set(),'sources':set(),'total_deals':0,'total_discount':0}
                    suppliers[tax_num]['names'].add(name)
                    suppliers[tax_num]['sources'].add('نموذج 41')
                    suppliers[tax_num]['total_deals']+=_sf(r.get('القيمة الإجمالية للتعامل',0))
                    suppliers[tax_num]['total_discount']+=_sf(r.get('المحصل لحساب الضريبة',0))
        vat_data=load_data(VAT_FILE)
        for rec in vat_data:
            for r in rec.get('records',[]):
                tax_num=str(r.get('رقم التسجيل الضريبي','')).strip()
                name=str(r.get('اسم الممول','')).strip()
                if tax_num and name:
                    if tax_num not in suppliers:
                        suppliers[tax_num]={'tax_number':tax_num,'names':set(),'sources':set(),'total_deals':0,'total_discount':0}
                    suppliers[tax_num]['names'].add(name)
                    suppliers[tax_num]['sources'].add('القيمة المضافة')
                    suppliers[tax_num]['total_deals']+=_sf(r.get('20% قيمة مضافة',0))
                    suppliers[tax_num]['total_discount']+=_sf(r.get('ضريبة الجدول',0))
        portal_out=load_data(PORTAL_OUT_FILE)
        for rec in portal_out:
            for r in rec.get('records',[]):
                tax_num=str(r.get('رقم التسجيل (المستلم)','') or r.get('رقم التسجيل (المصدر)','')).strip()
                name=str(r.get('اسم المستلم','') or r.get('اسم المصدر','')).strip()
                if not tax_num:
                    tax_num=str(r.get('رقم التسجيل (الطرف الآخر)','')).strip()
                if not name:
                    name=str(r.get('الطرف الآخر','')).strip()
                if tax_num and name:
                    if tax_num not in suppliers:
                        suppliers[tax_num]={'tax_number':tax_num,'names':set(),'sources':set(),'total_deals':0,'total_discount':0}
                    suppliers[tax_num]['names'].add(name)
                    suppliers[tax_num]['sources'].add('فواتير صادرة')
                    suppliers[tax_num]['total_deals']+=_sf(r.get('الإجمالي (بعد الضريبة)',0))
                    suppliers[tax_num]['total_discount']+=_sf(r.get('الخصم',0))
        portal_in=load_data(PORTAL_IN_FILE)
        for rec in portal_in:
            for r in rec.get('records',[]):
                tax_num=str(r.get('رقم التسجيل (المصدر)','') or r.get('رقم التسجيل (المستلم)','')).strip()
                name=str(r.get('اسم المصدر','') or r.get('اسم المستلم','')).strip()
                if not tax_num:
                    tax_num=str(r.get('رقم التسجيل (الطرف الآخر)','')).strip()
                if not name:
                    name=str(r.get('الطرف الآخر','')).strip()
                if tax_num and name:
                    if tax_num not in suppliers:
                        suppliers[tax_num]={'tax_number':tax_num,'names':set(),'sources':set(),'total_deals':0,'total_discount':0}
                    suppliers[tax_num]['names'].add(name)
                    suppliers[tax_num]['sources'].add('فواتير واردة')
                    suppliers[tax_num]['total_deals']+=_sf(r.get('الإجمالي (بعد الضريبة)',0))
                    suppliers[tax_num]['total_discount']+=_sf(r.get('الخصم',0))
        result=[]
        for tax_num,data in suppliers.items():
            data['names']=list(data['names'])
            data['sources']=list(data['sources'])
            result.append(data)
        return result

    suppliers_db=_build_suppliers_db()

    st.markdown(f"""<div style="padding:.6rem 1rem;border-radius:10px;background:rgba(116,185,255,.06);border:1px solid rgba(116,185,255,.15);color:#74b9ff;font-size:.82rem;margin-bottom:1rem;">
        📊 قاعدة البيانات تحتوي على <strong>{len(suppliers_db)}</strong> ممول من جميع المصادر
    </div>""",unsafe_allow_html=True)

    search_mode=st.radio("search_mode",["🔍 بحث بالسجل الضريبي","📝 بحث بالاسم"],horizontal=True,label_visibility="collapsed")

    if search_mode=="🔍 بحث بالسجل الضريبي":
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>بحث بالسجل الضريبي</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        tax_input=st.text_input("أدخل رقم التسجيل الضريبي",key="sup_tax_q",placeholder="مثال: 123-456-789")
        if st.button("🔍 بحث",key="sup_tax_btn",type="primary"):
            if not tax_input.strip():
                st.warning("أدخل رقم التسجيل الضريبي")
            else:
                query=tax_input.strip()
                results=[s for s in suppliers_db if query in s['tax_number']]
                if results:
                    st.success(f"تم العثور على {len(results)} ممول")
                    for s in results:
                        names_html=" • ".join([f'<span style="color:#a29bfe;">{n}</span>' for n in s['names']])
                        sources_html=" • ".join([f'<span style="background:rgba(108,92,231,.15);border:1px solid rgba(108,92,231,.2);border-radius:6px;padding:.15rem .5rem;color:#a29bfe;font-size:.7rem;">{src}</span>' for src in s['sources']])
                        st.markdown(f"""<div class="erp-card" style="margin-bottom:.8rem;border-left:3px solid #6c5ce7;">
                            <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                                <div>
                                    <div style="color:#fff;font-weight:700;font-size:1rem;margin-bottom:.3rem;">{names_html}</div>
                                    <div style="color:var(--text2);font-size:.82rem;margin-bottom:.4rem;">📋 رقم التسجيل: <strong style="color:#00cec9;">{s['tax_number']}</strong></div>
                                    <div style="margin-bottom:.3rem;">المصادر: {sources_html}</div>
                                </div>
                                <div style="text-align:left;">
                                    <div style="color:#55efc4;font-weight:700;font-size:.9rem;">{fmt(s['total_deals'])}</div>
                                    <div style="color:var(--text2);font-size:.7rem;">إجمالي التعاملات</div>
                                </div>
                            </div>
                        </div>""",unsafe_allow_html=True)
                else:
                    st.info(f"لم يتم العثور على ممول بالرقم: {query}")
        st.markdown('</div>',unsafe_allow_html=True)

    else:
        st.markdown('<div class="erp-section"><div class="erp-section-dot"></div><h3>بحث بالاسم</h3></div>',unsafe_allow_html=True)
        st.markdown('<div class="erp-card">',unsafe_allow_html=True)
        name_input=st.text_input("اكتب جزء من اسم الممول",key="sup_name_q",placeholder="مثال: شركة")
        if st.button("🔍 بحث",key="sup_name_btn",type="primary"):
            if not name_input.strip():
                st.warning("أدخل اسم الممول")
            else:
                query=name_input.strip()
                results=[]
                for s in suppliers_db:
                    for n in s['names']:
                        if query in n:
                            results.append(s)
                            break
                if results:
                    st.success(f"تم العثور على {len(results)} ممول")
                    for s in results:
                        names_html=" • ".join([f'<span style="color:#a29bfe;">{n}</span>' for n in s['names']])
                        sources_html=" • ".join([f'<span style="background:rgba(108,92,231,.15);border:1px solid rgba(108,92,231,.2);border-radius:6px;padding:.15rem .5rem;color:#a29bfe;font-size:.7rem;">{src}</span>' for src in s['sources']])
                        st.markdown(f"""<div class="erp-card" style="margin-bottom:.8rem;border-left:3px solid #6c5ce7;">
                            <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                                <div>
                                    <div style="color:#fff;font-weight:700;font-size:1rem;margin-bottom:.3rem;">{names_html}</div>
                                    <div style="color:var(--text2);font-size:.82rem;margin-bottom:.4rem;">📋 رقم التسجيل: <strong style="color:#00cec9;">{s['tax_number']}</strong></div>
                                    <div style="margin-bottom:.3rem;">المصادر: {sources_html}</div>
                                </div>
                                <div style="text-align:left;">
                                    <div style="color:#55efc4;font-weight:700;font-size:.9rem;">{fmt(s['total_deals'])}</div>
                                    <div style="color:var(--text2);font-size:.7rem;">إجمالي التعاملات</div>
                                </div>
                            </div>
                        </div>""",unsafe_allow_html=True)
                else:
                    st.info(f"لم يتم العثور على ممول بالاسم: {query}")
        st.markdown('</div>',unsafe_allow_html=True)

    if suppliers_db:
        st.markdown('<div class="erp-section" style="margin-top:1rem"><div class="erp-section-dot"></div><h3>جميع الممولين</h3></div>',unsafe_allow_html=True)
        sup_df=pd.DataFrame([{'رقم التسجيل الضريبي':s['tax_number'],'الاسم':' • '.join(s['names']),'المصادر':' • '.join(s['sources']),'إجمالي التعاملات':s['total_deals'],'إجمالي الخصومات':s['total_discount']} for s in suppliers_db])
        st.dataframe(sup_df,use_container_width=True,height=400)
        excel_buf=BytesIO()
        sup_df.to_excel(excel_buf,index=False,engine='xlsxwriter')
        excel_buf.seek(0)
        st.download_button("📊 تحميل قائمة الممولين",data=excel_buf.getvalue(),file_name="suppliers.xlsx",mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",key="dl_sup")
